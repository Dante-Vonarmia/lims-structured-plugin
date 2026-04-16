import io
import logging
import posixpath
import re
import zipfile
from copy import deepcopy
from datetime import datetime
from pathlib import Path
from typing import Any
from xml.etree import ElementTree as ET

from .docx_basis_utils import (
    extract_standard_code as _extract_standard_code,
    extract_standard_codes as _extract_standard_codes,
    extract_standard_codes_from_context_items as _extract_standard_codes_from_context_items,
    format_dual_mode_checkbox as _format_dual_mode_checkbox,
    infer_basis_mode as _infer_basis_mode,
    normalize_basis_mode as _normalize_basis_mode,
)
from .docx_cell_utils import (
    contains_compact_label as _contains_compact_label,
    find_cell_index_contains as _find_cell_index_contains,
    find_cell_index_with_text as _find_cell_index_with_text,
    get_cell_text,
    normalize_space,
    set_cell_text,
    split_date_parts,
)
from .docx_context_utils import (
    extract_basis_from_cell as _extract_basis_from_cell,
    resolve_report_dates as _resolve_report_dates,
    sanitize_context_date,
    sanitize_context_value,
    split_model_code_combined as _split_model_code_combined,
    tables_to_text_block as _tables_to_text_block,
)
from .docx_instrument_text_utils import (
    clean_item_name,
    extract_hammer_actual_rows,
    extract_hammer_actual_rows_from_context,
    extract_hammer_actual_rows_from_text,
    first_meaningful_value as _first_meaningful_value,
    looks_like_label as _looks_like_label,
    merge_hammer_actual_rows,
    merge_instrument_rows_with_catalog,
    normalize_catalog_token,
    parse_instrument_catalog_rows_json,
    parse_instrument_catalog_tokens,
    pick_cell as _pick_cell,
    sanitize_instrument_cell,
)
from .docx_xml_utils import (
    _capture_namespaces,
    _fill_page_number_placeholders_in_root,
    _fill_value_between_markers,
    _is_page_placeholder_text,
    _preserve_original_namespaces,
    _set_paragraph_text,
    set_cell_page_fields,
)
from .docx_semantic_bridge_utils import (
    _extract_humidity_from_other_calibration_info as _extract_humidity_from_other_calibration_info_bridge,
    _extract_location_from_other_calibration_info as _extract_location_from_other_calibration_info_bridge,
    _extract_section_measured_value as _extract_section_measured_value_bridge,
    _extract_section_uncertainty as _extract_section_uncertainty_bridge,
    _extract_temperature_from_other_calibration_info as _extract_temperature_from_other_calibration_info_bridge,
    _replace_measured_value as _replace_measured_value_bridge,
    _replace_uncertainty_value as _replace_uncertainty_value_bridge,
    _sanitize_location_text as _sanitize_location_text_bridge,
)
from .docx_media_dependency_utils import (
    _copy_docx_image_dependencies_for_nodes,
    _copy_docx_image_dependencies_for_table,
    _copy_docx_image_dependencies_for_tables,
)
from .docx_data_extraction_utils import (
    extract_any_date,
    extract_date_from_text,
    extract_docx_text,
    extract_instrument_rows,
    extract_value_by_regex,
    extract_value_from_tables,
    read_docx_tables,
)
from .extract_service import extract_fields
from .fixed_template_rule_engine import (
    fill_base_fields_in_cells_by_rules,
    fill_base_fields_in_paragraphs_by_rules,
    fill_base_fields_in_tables_by_rules,
    find_cell_index_contains_any,
    find_generic_record_table_by_rules,
)
from .r872_result_rules import fill_r872_requirement_text, should_mark_r872_result
from .result_check_matcher import extract_source_general_check_lines, match_best_source_line
from .semantic_fill_lib import (
    build_series_row_value_maps_from_general_check_text,
    build_semantic_value_maps_from_general_check_text,
    detect_semantic_key_value_columns,
    extract_humidity_from_other_calibration_info,
    extract_location_from_other_calibration_info,
    extract_measured_value_items,
    extract_section_measured_value,
    extract_section_uncertainty,
    extract_temperature_from_other_calibration_info,
    extract_text_block,
    extract_uncertainty_items,
    extract_uncertainty_u_value,
    is_detail_general_check_sparse,
    is_reliable_result_semantic_match,
    normalize_multiline_text,
    normalize_multiline_text_preserve_tabs,
    normalize_semantic_key,
    pick_series_row_values_for_label,
    pick_semantic_value_for_key,
    replace_measured_value,
    replace_uncertainty_value,
    replace_measured_value_placeholder_by_items,
    replace_uncertainty_u_placeholder_by_items,
    replace_uncertainty_u_placeholder,
    resolve_detail_general_check_for_generic_fill,
    sanitize_location_text,
    score_detail_general_check_text,
)
from .templates import fill_r846b_specific_sections
from .signature_store_file import resolve_signature_image_path
from .template_mapping_library_service import get_fill_placeholders

try:
    from docx import Document
    from docx.shared import Mm
except Exception:  # pragma: no cover - optional runtime dependency
    Document = None
    Mm = None

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
NS = {"w": W_NS, "r": R_NS}
DOC_XML_PATH = "word/document.xml"
DOC_RELS_PATH = "word/_rels/document.xml.rels"
CONTENT_TYPES_PATH = "[Content_Types].xml"
MAX_INSTRUMENT_ROWS = 5
_PLACEHOLDER_VALUES = {"", "-", "--", "—", "/", "／"}
logger = logging.getLogger(__name__)
_MODIFY_CERT_SIGNOFF_COMPANY = "成都金克星气体有限公司"


def fill_r801b_docx(
    template_path: Path,
    output_path: Path,
    context: dict[str, str],
    source_file_path: Path | None,
) -> bool:
    if not template_path.exists():
        return False

    payload = build_r801b_payload(context=context, source_file_path=source_file_path)
    if not payload:
        return False

    with zipfile.ZipFile(template_path, "r") as zin:
        original_xml = zin.read(DOC_XML_PATH)
        original_namespaces = _capture_namespaces(original_xml)
        root = ET.fromstring(original_xml)

    tables = root.findall(".//w:tbl", NS)
    if not tables:
        return False

    info_table = _find_info_table(tables)
    if info_table is not None:
        _fill_info_table(info_table, payload)

    record_table = _find_record_table(tables)
    if record_table is not None:
        _fill_record_table(record_table, payload)
    _fill_generic_base_labels_in_paragraphs(root, {"certificate_no": payload.get("certificate_no", "")})
    _fill_page_number_placeholders_in_root(root)

    _preserve_original_namespaces(root, original_namespaces)
    updated_xml = ET.tostring(root, encoding="utf-8", xml_declaration=True)
    with zipfile.ZipFile(template_path, "r") as zin, zipfile.ZipFile(output_path, "w") as zout:
        for item in zin.infolist():
            if item.filename == DOC_XML_PATH:
                zout.writestr(item, updated_xml)
            elif _is_header_xml_path(item.filename):
                header_xml = _fill_header_base_fields_xml(
                    zin.read(item.filename),
                    {"certificate_no": payload.get("certificate_no", "")},
                )
                zout.writestr(item, header_xml)
            else:
                zout.writestr(item, zin.read(item.filename))
    return True


def fill_r803b_docx(
    template_path: Path,
    output_path: Path,
    context: dict[str, str],
    source_file_path: Path | None,
) -> bool:
    if not template_path.exists():
        return False

    payload = build_r803b_payload(context=context, source_file_path=source_file_path)
    if not payload:
        return False

    with zipfile.ZipFile(template_path, "r") as zin:
        original_xml = zin.read(DOC_XML_PATH)
        original_namespaces = _capture_namespaces(original_xml)
        root = ET.fromstring(original_xml)

    tables = root.findall(".//w:tbl", NS)
    if not tables:
        return False

    record_table = _find_r803b_record_table(tables)
    if record_table is None:
        return False

    _fill_r803b_record_table(record_table, payload)
    _fill_r803b_result_checks(record_table)
    _fill_page_number_placeholders_in_root(root)

    _preserve_original_namespaces(root, original_namespaces)
    updated_xml = ET.tostring(root, encoding="utf-8", xml_declaration=True)
    with zipfile.ZipFile(template_path, "r") as zin, zipfile.ZipFile(output_path, "w") as zout:
        for item in zin.infolist():
            if item.filename == DOC_XML_PATH:
                zout.writestr(item, updated_xml)
            elif _is_header_xml_path(item.filename):
                header_xml = _fill_header_base_fields_xml(zin.read(item.filename), payload)
                zout.writestr(item, header_xml)
            else:
                zout.writestr(item, zin.read(item.filename))
    return True


def fill_r802b_docx(
    template_path: Path,
    output_path: Path,
    context: dict[str, str],
    source_file_path: Path | None,
) -> bool:
    if not template_path.exists():
        return False

    payload = build_r802b_payload(context=context, source_file_path=source_file_path)
    if not payload:
        return False

    with zipfile.ZipFile(template_path, "r") as zin:
        original_xml = zin.read(DOC_XML_PATH)
        original_namespaces = _capture_namespaces(original_xml)
        root = ET.fromstring(original_xml)

    tables = root.findall(".//w:tbl", NS)
    if not tables:
        return False

    record_table = _find_r802b_record_table(tables)
    if record_table is not None:
        _fill_r802b_record_table(record_table, payload)
    _fill_page_number_placeholders_in_root(root)
    copied_general_check_table, copied_nodes = _copy_r802b_general_check_table_from_source(root, source_file_path)
    rel_updates: dict[str, bytes] = {}
    if copied_general_check_table and copied_nodes and source_file_path is not None:
        rel_updates = _copy_docx_image_dependencies_for_nodes(
            template_path=template_path,
            source_file_path=source_file_path,
            nodes=copied_nodes,
        )
    if not copied_general_check_table:
        _append_r802b_general_check_text_only(root, payload)

    _preserve_original_namespaces(root, original_namespaces)
    updated_xml = ET.tostring(root, encoding="utf-8", xml_declaration=True)
    with zipfile.ZipFile(template_path, "r") as zin, zipfile.ZipFile(output_path, "w") as zout:
        updated_rels_xml = rel_updates.get("__rels__")
        updated_ct_xml = rel_updates.get("__ct__")
        existing_names = {x.filename for x in zin.infolist()}
        for item in zin.infolist():
            if item.filename == DOC_XML_PATH:
                zout.writestr(item, updated_xml)
            elif _is_header_xml_path(item.filename):
                header_xml = _fill_header_base_fields_xml(zin.read(item.filename), payload)
                zout.writestr(item, header_xml)
            elif item.filename == DOC_RELS_PATH and updated_rels_xml is not None:
                zout.writestr(item, updated_rels_xml)
            elif item.filename == CONTENT_TYPES_PATH and updated_ct_xml is not None:
                zout.writestr(item, updated_ct_xml)
            else:
                zout.writestr(item, zin.read(item.filename))
        if updated_rels_xml is not None and DOC_RELS_PATH not in existing_names:
            zout.writestr(DOC_RELS_PATH, updated_rels_xml)
        if updated_ct_xml is not None and CONTENT_TYPES_PATH not in existing_names:
            zout.writestr(CONTENT_TYPES_PATH, updated_ct_xml)
        for path, raw in rel_updates.items():
            if path.startswith("__"):
                continue
            zout.writestr(path, raw)
    return True


def fill_r825b_docx(
    template_path: Path,
    output_path: Path,
    context: dict[str, str],
    source_file_path: Path | None,
) -> bool:
    if not template_path.exists():
        return False

    payload = build_r825b_payload(context=context, source_file_path=source_file_path)
    if not payload:
        return False

    with zipfile.ZipFile(template_path, "r") as zin:
        original_xml = zin.read(DOC_XML_PATH)
        original_namespaces = _capture_namespaces(original_xml)
        root = ET.fromstring(original_xml)

    tables = root.findall(".//w:tbl", NS)
    if not tables:
        return False

    record_table = _find_r825b_record_table(tables)
    if record_table is None:
        return False

    _fill_r825b_record_table(record_table, payload)
    _fill_page_number_placeholders_in_root(root)
    _strip_general_check_required_marker(root)

    _preserve_original_namespaces(root, original_namespaces)
    updated_xml = ET.tostring(root, encoding="utf-8", xml_declaration=True)
    with zipfile.ZipFile(template_path, "r") as zin, zipfile.ZipFile(output_path, "w") as zout:
        for item in zin.infolist():
            if item.filename == DOC_XML_PATH:
                zout.writestr(item, updated_xml)
            elif _is_header_xml_path(item.filename):
                header_xml = _fill_header_base_fields_xml(zin.read(item.filename), payload)
                zout.writestr(item, header_xml)
            else:
                zout.writestr(item, zin.read(item.filename))
    return True


def fill_r846b_docx(
    template_path: Path,
    output_path: Path,
    context: dict[str, str],
    source_file_path: Path | None,
) -> bool:
    if not template_path.exists():
        return False

    payload = build_r825b_payload(context=context, source_file_path=source_file_path)
    if not payload:
        return False

    detail_general_check = normalize_multiline_text_preserve_tabs(
        context.get("general_check_full", ""),
        normalize_space=normalize_space,
    ) or normalize_multiline_text(context.get("general_check", ""), normalize_space=normalize_space)
    raw_record_text = normalize_multiline_text(context.get("raw_record", ""), normalize_space=normalize_space)
    source_text = detail_general_check or raw_record_text
    if detail_general_check:
        payload["detail_general_check"] = detail_general_check
    payload["__raw_record_text"] = raw_record_text
    payload["__template_name"] = template_path.name

    with zipfile.ZipFile(template_path, "r") as zin:
        original_xml = zin.read(DOC_XML_PATH)
        original_namespaces = _capture_namespaces(original_xml)
        root = ET.fromstring(original_xml)

    tables = root.findall(".//w:tbl", NS)
    changed = False
    if tables:
        record_table = _find_generic_record_table(tables)
        if record_table is not None:
            _fill_r825b_record_table(record_table, payload)
            changed = True
        else:
            changed = _fill_generic_base_labels_in_tables(tables, payload)
        changed = fill_r846b_specific_sections(
            tables,
            source_text,
            ns=NS,
            placeholder_values=_PLACEHOLDER_VALUES,
            normalize_space=normalize_space,
            get_cell_text=get_cell_text,
            set_cell_text=set_cell_text,
            extract_value_by_regex=extract_value_by_regex,
            replace_uncertainty_value=_replace_uncertainty_value,
        ) or changed
    if extract_source_general_check_lines(str(payload.get("detail_general_check", ""))):
        changed = _fill_generic_result_checks_by_semantics(tables, payload) or changed
    _fill_page_number_placeholders_in_root(root)
    changed = _fill_generic_base_labels_in_paragraphs(root, payload) or changed
    changed = _strip_general_check_required_marker(root) or changed

    _preserve_original_namespaces(root, original_namespaces)
    updated_xml = ET.tostring(root, encoding="utf-8", xml_declaration=True)
    with zipfile.ZipFile(template_path, "r") as zin, zipfile.ZipFile(output_path, "w") as zout:
        for item in zin.infolist():
            if item.filename == DOC_XML_PATH:
                zout.writestr(item, updated_xml)
            elif _is_header_xml_path(item.filename):
                header_xml = _fill_header_base_fields_xml(zin.read(item.filename), payload)
                zout.writestr(item, header_xml)
            else:
                zout.writestr(item, zin.read(item.filename))
    return True


def fill_generic_record_docx(
    template_path: Path,
    output_path: Path,
    context: dict[str, str],
    source_file_path: Path | None,
) -> bool:
    if not template_path.exists():
        return False

    payload = build_r825b_payload(context=context, source_file_path=source_file_path)
    if not payload:
        return False
    detail_general_check = _resolve_detail_general_check_for_generic_fill(context)
    detail_general_check = _trim_general_check_note_block_for_record_fill(detail_general_check)
    if detail_general_check:
        payload["detail_general_check"] = detail_general_check
    payload["__raw_record_text"] = normalize_multiline_text(context.get("raw_record", ""), normalize_space=normalize_space)
    payload["__template_name"] = template_path.name
    payload["__measurement_items_text"] = normalize_multiline_text_preserve_tabs(
        context.get("measurement_items", ""),
        normalize_space=normalize_space,
    )
    for key in (
        "shield_background_noise_0kv_pc",
        "shield_background_noise_working_kv_pc",
        "shield_p1_dbm_values",
        "shield_p2_dbm_values",
        "shield_se_db_values",
        "shield_se_avg_db",
    ):
        payload[key] = normalize_space(str(context.get(key, "") or ""))

    with zipfile.ZipFile(template_path, "r") as zin:
        original_xml = zin.read(DOC_XML_PATH)
        original_namespaces = _capture_namespaces(original_xml)
        root = ET.fromstring(original_xml)

    tables = root.findall(".//w:tbl", NS)
    changed = False
    if tables:
        record_table = _find_generic_record_table(tables)
        if record_table is not None:
            _fill_r825b_record_table(record_table, payload)
            changed = True
        changed = _fill_generic_base_labels_in_tables(tables, payload) or changed
        changed = _fill_r882b_specific_sections(root, tables, payload) or changed
        changed = _fill_generic_semantic_value_matrices_from_payload(tables, payload) or changed
        changed = _fill_generic_semantic_series_rows_from_payload(tables, payload) or changed
        changed = _fill_generic_uncertainty_u_from_payload(root, tables, payload) or changed
    if extract_source_general_check_lines(str(payload.get("detail_general_check", ""))):
        changed = _fill_generic_result_checks_by_semantics(tables, payload) or changed
    elif _should_auto_fill_result_checks(template_path=template_path, payload=payload):
        changed = _fill_generic_result_checks_in_tables(tables) or changed
    _fill_page_number_placeholders_in_root(root)
    placeholder_values = _build_placeholder_values_from_rules(
        template_name=template_path.name,
        context=context,
        payload=payload,
    )
    _fill_text_placeholders_in_root(root, placeholder_values)
    _fill_jinja_placeholders_in_root(
        root,
        _build_jinja_placeholder_values(context=context, payload=payload),
    )
    _fill_appendix1_table_rows_from_context(root, context)
    changed = _fill_generic_base_labels_in_paragraphs(root, payload) or changed
    changed = _strip_general_check_required_marker(root) or changed

    _write_docx_with_updated_root(
        template_path=template_path,
        output_path=output_path,
        root=root,
        original_namespaces=original_namespaces,
        header_payload=payload,
    )
    _fill_signature_images_for_output_docx(output_path, payload)
    return True


def _trim_general_check_note_block_for_record_fill(text: str) -> str:
    value = normalize_multiline_text_preserve_tabs(text or "", normalize_space=normalize_space)
    if not value:
        return ""
    lines: list[str] = []
    for raw in value.splitlines():
        line = str(raw or "").strip()
        if re.search(r"^注[:：]?|^Notes?[:：]?", line, flags=re.IGNORECASE):
            break
        lines.append(raw)
    return normalize_multiline_text_preserve_tabs("\n".join(lines), normalize_space=normalize_space)


def _strip_general_check_required_marker(root: ET.Element) -> bool:
    changed = False
    for node in root.findall(".//w:t", NS):
        original = str(node.text or "")
        updated = re.sub(r"(一般检查)\s*[（(]\s*\*\s*[）)]", r"\1", original)
        updated = re.sub(r"[（(]\s*\*\s*[）)]", "", updated)
        if updated != original:
            node.text = updated
            changed = True
    return changed


def _normalize_modify_certificate_standard_code(value: Any) -> str:
    text = normalize_space(str(value or ""))
    if not text:
        return ""
    code = _extract_standard_code(text)
    if code:
        return str(code).upper().replace(" ", "")
    compact = text.upper().replace(" ", "")
    if "13004" in compact:
        return "GB/T13004-2016"
    if "13077" in compact:
        return "GB/T13077-2024"
    return ""


def _fill_modify_certificate_standard_checkboxes(
    root: ET.Element,
    payload: dict[str, Any],
    context: dict[str, Any],
) -> bool:
    standard_value = normalize_space(
        context.get("inspect_standard", ""),
    ) or normalize_space(payload.get("inspect_standard", ""))
    standard_code = _normalize_modify_certificate_standard_code(standard_value)
    if not standard_code:
        return False

    changed = False
    checked_char = "00FE"
    unchecked_char = "00A8"
    paragraph_rules = (
        ("钢质无缝气瓶定期检验与评定", standard_code.startswith("GB/T13004")),
        ("铝合金无缝气瓶定期检验与评定", standard_code.startswith("GB/T13077")),
    )

    for paragraph in root.findall(".//w:p", NS):
        text = normalize_space("".join(str(node.text or "") for node in paragraph.findall(".//w:t", NS)))
        if not text:
            continue
        for marker, should_check in paragraph_rules:
            if marker not in text:
                continue
            symbols = paragraph.findall(".//w:sym", NS)
            if not symbols:
                break
            target_char = checked_char if should_check else unchecked_char
            current_char = str(symbols[0].attrib.get(f"{{{W_NS}}}char", "") or "").strip()
            if current_char != target_char:
                symbols[0].set(f"{{{W_NS}}}char", target_char)
                changed = True
            break
    return changed


def _resolve_modify_certificate_signoff_date(context: dict[str, Any], key: str) -> str:
    direct_value = sanitize_context_date(str(context.get(key, "") or ""))
    if direct_value:
        return direct_value
    for fallback_key in ("report_date", "release_date", "publish_date", "calibration_date", "receive_date"):
        candidate = sanitize_context_date(str(context.get(fallback_key, "") or ""))
        if candidate:
            return candidate
    return ""


def _replace_modify_certificate_date_text(text: str, date_value: str) -> str:
    if not date_value:
        return text
    return re.sub(
        r"(日期\s*[:：]\s*)(?:\d{4}年\d{1,2}月\d{1,2}日|【[^】]+】|_{2,})",
        lambda m: f"{m.group(1)}{date_value}",
        str(text or ""),
        count=1,
    )


def _replace_first_signoff_date_in_nodes(text_nodes: list[ET.Element], date_value: str) -> bool:
    if not date_value:
        return False
    for node in text_nodes:
        original = str(node.text or "")
        if not original:
            continue
        updated = _replace_modify_certificate_date_text(original, date_value)
        if updated != original:
            node.text = updated
            return True
    return False


def _replace_signoff_company_in_nodes(text_nodes: list[ET.Element], company_name: str) -> bool:
    if not company_name:
        return False
    changed = False
    for node in text_nodes:
        original = str(node.text or "")
        if not original or "有限公司" not in original:
            continue
        updated = re.sub(
            r"[^\s：:]{1,120}有限公司",
            company_name,
            original,
            count=1,
        )
        if updated != original:
            node.text = updated
            changed = True
            break
    return changed


def _fill_modify_certificate_signoff_sections(root: ET.Element, context: dict[str, Any]) -> bool:
    inspector_date = _resolve_modify_certificate_signoff_date(context, "inspector_sign_date")
    reviewer_date = _resolve_modify_certificate_signoff_date(context, "reviewer_sign_date")
    approver_date = _resolve_modify_certificate_signoff_date(context, "approver_sign_date")
    company_sign_date = _resolve_modify_certificate_signoff_date(context, "company_sign_date")

    changed = False
    for paragraph in root.findall(".//w:p", NS):
        text_nodes = paragraph.findall(".//w:t", NS)
        if not text_nodes:
            continue

        paragraph_text = "".join(str(node.text or "") for node in text_nodes)
        compact = normalize_space(paragraph_text)
        if not compact:
            continue

        if "检验员（签字）" in compact and inspector_date:
            if _replace_first_signoff_date_in_nodes(text_nodes, inspector_date):
                changed = True
            continue

        if "审核（签字）" in compact:
            if reviewer_date and _replace_first_signoff_date_in_nodes(text_nodes, reviewer_date):
                changed = True
            if approver_date and _replace_first_signoff_date_in_nodes(text_nodes, approver_date):
                changed = True
            continue

        if "有限公司" in compact and "报告抬头公司" not in compact:
            if _replace_signoff_company_in_nodes(text_nodes, _MODIFY_CERT_SIGNOFF_COMPANY):
                changed = True
            if company_sign_date and _replace_first_signoff_date_in_nodes(text_nodes, company_sign_date):
                changed = True
            continue

    return changed


def fill_modify_certificate_docx(
    template_path: Path,
    output_path: Path,
    context: dict[str, str],
    source_file_path: Path | None,
) -> bool:
    if not template_path.exists():
        return False

    payload = build_r825b_payload(context=context, source_file_path=source_file_path)
    if not payload:
        return False
    detail_general_check = _resolve_detail_general_check_for_generic_fill(context)
    if detail_general_check:
        payload["detail_general_check"] = detail_general_check

    with zipfile.ZipFile(template_path, "r") as zin:
        original_xml = zin.read(DOC_XML_PATH)
        original_namespaces = _capture_namespaces(original_xml)
        root = ET.fromstring(original_xml)

    rel_updates: dict[str, bytes] = {}
    tables = root.findall(".//w:tbl", NS)
    changed = False
    if tables:
        changed = _fill_modify_certificate_front_page_sections(tables, payload) or changed
        changed = _fill_modify_certificate_blueprint_sections(tables, payload, context) or changed
    copied_continued_page_table, _copied_tables = _copy_modify_certificate_continued_page_table_from_source(
        target_root=root,
        source_file_path=source_file_path,
    )
    if copied_continued_page_table and source_file_path is not None:
        body = root.find("./w:body", NS)
        if body is not None:
            range_idx = _find_modify_certificate_continued_body_range(body)
            if range_idx is not None:
                start, end = range_idx
                copied_nodes = list(body)[start:end]
                rel_updates = _copy_docx_image_dependencies_for_nodes(
                    template_path=template_path,
                    source_file_path=source_file_path,
                    nodes=copied_nodes,
                )
        changed = True
    _fill_page_number_placeholders_in_root(root)
    changed = _fill_modify_certificate_standard_checkboxes(root, payload, context) or changed
    changed = _fill_appendix1_table_rows_from_context(root, context) or changed
    changed = _fill_modify_certificate_signoff_sections(root, context) or changed
    changed = _fill_generic_base_labels_in_paragraphs(
        root,
        {"certificate_no": payload.get("certificate_no", "")},
    ) or changed

    _write_docx_with_updated_root(
        template_path=template_path,
        output_path=output_path,
        root=root,
        original_namespaces=original_namespaces,
        header_payload={"certificate_no": payload.get("certificate_no", "")},
        rel_updates=rel_updates,
    )
    return True


def _fill_modify_certificate_front_page_sections(
    tables: list[ET.Element],
    payload: dict[str, Any],
) -> bool:
    value_mappings: tuple[tuple[str, tuple[str, ...]], ...] = (
        ("certificate_no", (r"缆\s*专\s*检\s*号", r"Certificate\s*series\s*number")),
        ("client_name", (r"委\s*托\s*单\s*位", r"Client")),
        ("address", (r"地\s*址", r"Address")),
        ("device_name", (r"器\s*具\s*名\s*称", r"Instrument\s*name")),
        ("manufacturer", (r"制\s*造\s*厂\s*/\s*商", r"Manufacturer")),
        ("device_model", (r"型\s*号\s*/\s*规\s*格", r"Model\s*/\s*Specification")),
        ("device_code", (r"器\s*具\s*编\s*号", r"Instrument\s*serial\s*number")),
    )
    date_mappings: tuple[tuple[str, str], ...] = (("发布日期", "publish_date"),)

    changed = False
    for tbl in tables:
        table_text = normalize_space(" ".join([get_cell_text(tc) for tc in tbl.findall(".//w:tc", NS)]))
        if not table_text:
            continue
        if "CALIBRATION CERTIFICATE" not in table_text:
            continue
        if "委托单位" not in table_text or "器具名称" not in table_text:
            continue
        rows = tbl.findall("./w:tr", NS)
        for row in rows:
            cells = row.findall("./w:tc", NS)
            if not cells:
                continue
            for label, payload_key in date_mappings:
                changed = _fill_modify_certificate_split_date_cell(
                    cells,
                    label=label,
                    date_text=str(payload.get(payload_key, "") or ""),
                ) or changed
            for payload_key, patterns in value_mappings:
                changed = _fill_modify_certificate_value_cell(
                    cells,
                    patterns=patterns,
                    value=normalize_space(payload.get(payload_key, "")),
                ) or changed
        break
    return changed


def _fill_modify_certificate_value_cell(
    cells: list[ET.Element],
    patterns: tuple[str, ...],
    value: str,
) -> bool:
    if not value:
        return False
    for idx, cell in enumerate(cells):
        text = normalize_space(get_cell_text(cell))
        if not text:
            continue
        if not any(re.search(pattern, text, flags=re.IGNORECASE) for pattern in patterns):
            continue
        target_idx = idx + 1 if idx + 1 < len(cells) else idx
        if target_idx == idx:
            return False
        set_cell_text(cells[target_idx], value)
        return True
    return False


def _fill_modify_certificate_split_date_cell(
    cells: list[ET.Element],
    label: str,
    date_text: str,
) -> bool:
    date_value = sanitize_context_date(date_text)
    if not date_value:
        return False
    if not any(_contains_compact_label(get_cell_text(cell), label) for cell in cells):
        return False
    parts = split_date_parts(date_value)
    if not parts:
        return False
    year, month, day = parts
    _fill_split_date(cells, label, f"{year}年{month}月{day}日")
    return True


def _fill_modify_certificate_blueprint_sections(
    tables: list[ET.Element],
    payload: dict[str, Any],
    context: dict[str, str],
) -> bool:
    changed = False
    for tbl in tables:
        table_text = normalize_space(" ".join([get_cell_text(tc) for tc in tbl.findall(".//w:tc", NS)]))
        if not table_text:
            continue
        if re.search(r"缆\s*专\s*检\s*号|Certificate\s*series\s*number", table_text, flags=re.IGNORECASE):
            changed = _fill_modify_certificate_continued_certificate_no(tbl, payload, context) or changed
        if (
            ("Main measurement standard instruments" in table_text or "主要计量标准器具" in table_text)
            and ("Calibration Information" in table_text or "校准信息" in table_text)
            and ("Received date" in table_text or "收样日期" in table_text)
        ):
            changed = _fill_modify_certificate_middle_table(tbl, payload, context) or changed
        if re.search(r"校准结果\s*/\s*说明|Results of calibration and additional explanation", table_text, flags=re.IGNORECASE):
            changed = _fill_modify_certificate_general_check_table(tbl, payload, context) or changed
    return changed


def _fill_modify_certificate_middle_table(
    tbl: ET.Element,
    payload: dict[str, Any],
    context: dict[str, str],
) -> bool:
    rows = tbl.findall("./w:tr", NS)
    if len(rows) < 4:
        return False
    changed = False

    changed = _fill_modify_certificate_continued_certificate_no(tbl, payload, context) or changed

    basis_lines = _resolve_basis_lines_for_blueprint(payload, context)
    if basis_lines:
        changed = _fill_modify_certificate_basis_rows(tbl, basis_lines) or changed

    measurement_rows = _resolve_measurement_rows_for_blueprint(context)
    changed = _fill_modify_certificate_measurement_rows(tbl, measurement_rows) or changed

    changed = _fill_modify_certificate_calibration_info_rows(tbl, payload, context) or changed

    return changed


def _fill_modify_certificate_continued_certificate_no(
    tbl: ET.Element,
    payload: dict[str, Any],
    context: dict[str, str],
) -> bool:
    certificate_no = normalize_space(payload.get("certificate_no", "")) or normalize_space(context.get("certificate_no", ""))
    if not certificate_no:
        return False

    changed = False
    rows = tbl.findall("./w:tr", NS)
    for row in rows:
        cells = row.findall("./w:tc", NS)
        if not cells:
            continue
        for idx, cell in enumerate(cells):
            text = normalize_space(get_cell_text(cell))
            if not text:
                continue
            if not re.search(r"缆\s*专\s*检\s*号|Certificate\s*series\s*number", text, flags=re.IGNORECASE):
                continue
            next_idx = idx + 1
            if next_idx < len(cells):
                next_text = normalize_space(get_cell_text(cells[next_idx]))
                if not re.search(r"缆\s*专\s*检\s*号|Certificate\s*series\s*number", next_text, flags=re.IGNORECASE):
                    set_cell_text(cells[next_idx], certificate_no)
                    changed = True
                    break
            set_cell_text(cells[idx], f"缆专检号：{certificate_no}")
            changed = True
            break
    return changed


_MODIFY_CERT_MEASUREMENT_COL_RULES: tuple[tuple[str, tuple[str, ...], int], ...] = (
    ("name", ("器具名称", "Instrument name"), 0),
    ("model", ("型号/规格", "Model/Specification"), 1),
    ("code", ("器具编号", "仪器编号", "设备编号", "编号"), 2),
    ("range", ("测量范围", "Measurement range"), 3),
    ("uncertainty", ("不确定度", "Accuracy", "Uncertainty"), 4),
    ("cert_valid", ("证书编号", "有效期限", "Certificate number", "Valid date"), 5),
    ("trace", ("溯源机构", "traceability"), 6),
)


def _modify_cert_row_text(row: ET.Element) -> str:
    return normalize_space(" ".join([get_cell_text(tc) for tc in row.findall("./w:tc", NS)]))


def _find_first_row_index(rows: list[ET.Element], pattern: str) -> int:
    return next((idx for idx, row in enumerate(rows) if re.search(pattern, _modify_cert_row_text(row), flags=re.IGNORECASE)), -1)


def _expand_data_rows_before_anchor(tbl: ET.Element, data_start: int, data_end: int, need_rows: int) -> int:
    existing_rows = data_end - data_start
    if existing_rows <= 0 or need_rows <= existing_rows:
        return existing_rows
    rows = tbl.findall("./w:tr", NS)
    template_row = rows[data_start]
    for _ in range(need_rows - existing_rows):
        rows = tbl.findall("./w:tr", NS)
        anchor_row = rows[data_end]
        anchor_pos = list(tbl).index(anchor_row)
        tbl.insert(anchor_pos, deepcopy(template_row))
        data_end += 1
    rows = tbl.findall("./w:tr", NS)
    return data_end - data_start


def _fill_modify_certificate_basis_rows(tbl: ET.Element, basis_lines: list[str]) -> bool:
    rows = tbl.findall("./w:tr", NS)
    basis_title_idx = _find_first_row_index(rows, r"本次校准所依据的技术规范|Reference documents for the calibration")
    measurement_title_idx = _find_first_row_index(rows, r"本次校准所使用的主要计量标准器具|Main measurement standard instruments")
    if basis_title_idx < 0 or measurement_title_idx <= basis_title_idx + 1:
        return False
    basis_start = basis_title_idx + 1
    basis_end = measurement_title_idx
    existing_rows = _expand_data_rows_before_anchor(tbl, basis_start, basis_end, len(basis_lines))
    rows = tbl.findall("./w:tr", NS)
    changed = False
    for offset in range(max(existing_rows, 0)):
        cells = rows[basis_start + offset].findall("./w:tc", NS)
        if not cells:
            continue
        target_idx = 1 if len(cells) > 1 else 0
        value = basis_lines[offset] if offset < len(basis_lines) else ""
        set_cell_text(cells[target_idx], value)
        changed = True
    return changed


def _resolve_modify_cert_measurement_col_map(header_cells: list[ET.Element]) -> dict[str, int]:
    col_map: dict[str, int] = {}
    for key, markers, fallback in _MODIFY_CERT_MEASUREMENT_COL_RULES:
        idx = _find_cell_index_contains_any(header_cells, markers)
        col_map[key] = idx if idx >= 0 else fallback
    if 0 <= col_map["code"] < len(header_cells):
        if "证书" in get_cell_text(header_cells[col_map["code"]]):
            col_map["code"] = 2 if len(header_cells) > 2 else -1
    return col_map


def _fill_modify_certificate_measurement_rows(tbl: ET.Element, measurement_rows: list[dict[str, str]]) -> bool:
    rows = tbl.findall("./w:tr", NS)
    header_idx = next(
        (
            idx
            for idx, row in enumerate(rows)
            if re.search(r"器具名称|Instrument name", _modify_cert_row_text(row), flags=re.IGNORECASE)
            and re.search(r"测量范围|Measurement range", _modify_cert_row_text(row), flags=re.IGNORECASE)
        ),
        -1,
    )
    summary_idx = _find_first_row_index(rows, r"以上计量标准器具|Quantity values of above measurement standards")
    if header_idx < 0 or summary_idx <= header_idx + 1:
        return False
    data_start = header_idx + 1
    data_end = summary_idx
    existing_rows = _expand_data_rows_before_anchor(tbl, data_start, data_end, len(measurement_rows))
    rows = tbl.findall("./w:tr", NS)

    changed = False
    header_cells = rows[header_idx].findall("./w:tc", NS)
    if len(header_cells) > 2 and not normalize_space(get_cell_text(header_cells[2])):
        set_cell_text(header_cells[2], "编号\nNumber")
        changed = True
    col_map = _resolve_modify_cert_measurement_col_map(header_cells)
    for offset in range(max(existing_rows, 0)):
        row_cells = rows[data_start + offset].findall("./w:tc", NS)
        item = measurement_rows[offset] if offset < len(measurement_rows) else {}
        for key, idx in col_map.items():
            if idx < 0 or idx >= len(row_cells):
                continue
            set_cell_text(row_cells[idx], normalize_space(str(item.get(key, ""))))
            changed = True
    return changed


def _fill_modify_certificate_calibration_info_rows(
    tbl: ET.Element,
    payload: dict[str, Any],
    context: dict[str, str],
) -> bool:
    location = normalize_space(payload.get("location", "")) or normalize_space(context.get("location", ""))
    temp = normalize_space(payload.get("temperature", "")) or normalize_space(context.get("temperature", ""))
    humidity = normalize_space(payload.get("humidity", "")) or normalize_space(context.get("humidity", ""))
    other = normalize_space(context.get("calibration_other", ""))
    receive_date = sanitize_context_date(context.get("receive_date", ""))
    calibration_date = sanitize_context_date(context.get("calibration_date", ""))

    rows = tbl.findall("./w:tr", NS)
    changed = False
    for row in rows:
        cells = row.findall("./w:tc", NS)
        if not cells:
            continue
        text = _modify_cert_row_text(row)
        if location and re.search(r"地点", text) and not re.search(r"收样日期|校准日期", text):
            set_cell_text(cells[-1], f"地点： {location}\n")
            changed = True
        if re.search(r"温度", text) and re.search(r"湿度", text):
            if temp:
                idx = _find_cell_index_contains(cells, "温度")
                if idx >= 0:
                    set_cell_text(cells[idx], f"温度：{temp}\nAmbient temperature")
                    changed = True
            if humidity:
                idx = _find_cell_index_contains(cells, "湿度")
                if idx >= 0:
                    set_cell_text(cells[idx], f"湿度：{humidity}\nRelative humidity")
                    changed = True
            if other:
                idx = _find_cell_index_contains_any(cells, ("其它", "其他", "Others"))
                if idx >= 0:
                    other_value = "/" if other in {"/", "／"} else other
                    set_cell_text(cells[idx], f"其它： {other_value}\nOthers")
                    changed = True
        if re.search(r"收\s*样\s*日\s*期", text) and re.search(r"校\s*准\s*日\s*期", text):
            if receive_date:
                _fill_split_date(cells, "收样日期", receive_date)
                changed = True
            if calibration_date:
                _fill_split_date(cells, "校准日期", calibration_date)
                changed = True
    return changed


def _fill_modify_certificate_general_check_table(
    tbl: ET.Element,
    payload: dict[str, Any],
    context: dict[str, str],
) -> bool:
    rows = tbl.findall("./w:tr", NS)
    if len(rows) < 2:
        return False
    source = normalize_multiline_text_preserve_tabs(
        context.get("general_check_full", ""),
        normalize_space=normalize_space,
    ) or normalize_multiline_text(
        payload.get("detail_general_check", ""),
        normalize_space=normalize_space,
    )
    lines = _normalize_general_check_lines_for_blueprint(source)
    if not lines:
        return False
    row1_cells = rows[1].findall("./w:tc", NS)
    if not row1_cells:
        return False
    set_cell_text(row1_cells[0], "\n".join(lines))
    return True


def _resolve_basis_lines_for_blueprint(payload: dict[str, Any], context: dict[str, str]) -> list[str]:
    raw_items = context.get("basis_standard_items", "")
    items: list[str] = []
    if isinstance(raw_items, list):
        items = [normalize_space(str(item or "")) for item in raw_items]
    elif isinstance(raw_items, tuple):
        items = [normalize_space(str(item or "")) for item in raw_items]
    else:
        text_items = normalize_multiline_text_preserve_tabs(str(raw_items or ""), normalize_space=normalize_space)
        if text_items:
            items = [normalize_space(part) for part in re.split(r"[\n]+", text_items)]
    items = [item for item in items if item]
    if items:
        deduped: list[str] = []
        seen: set[str] = set()
        for item in items:
            if item in seen:
                continue
            seen.add(item)
            deduped.append(item)
        return deduped
    basis_text = normalize_space(context.get("basis_standard", "")) or normalize_space(payload.get("basis_standard", ""))
    if basis_text:
        lines = [normalize_space(line) for line in str(basis_text).splitlines() if normalize_space(line)]
        if lines:
            return lines
        return [basis_text]
    code_fallback = _extract_standard_codes_from_context_items(context.get("basis_standard_items", "")) or _extract_standard_codes(
        normalize_space(payload.get("basis_standard", "")),
    )
    if code_fallback:
        return code_fallback
    return []


def _parse_measurement_items_rows(text: str) -> list[dict[str, str]]:
    source = normalize_multiline_text_preserve_tabs(text, normalize_space=normalize_space)
    if not source:
        return []
    lines = [line for line in source.splitlines() if normalize_space(line)]
    if not lines:
        return []
    rows: list[list[str]] = []
    for line in lines:
        parts = [normalize_space(part) for part in line.split("\t")]
        if len(parts) < 2:
            parts = [normalize_space(part) for part in re.split(r"\s{2,}", line)]
        if len(parts) < 2 and "|" in line:
            parts = [normalize_space(part) for part in line.split("|")]
        parts = [p for p in parts if p]
        if not parts:
            continue
        rows.append(parts)
    if not rows:
        return []
    header_tokens = "".join(rows[0]).lower()
    if ("器具名称" in rows[0][0]) or ("instrumentname" in re.sub(r"\s+", "", header_tokens)):
        rows = rows[1:]
    result: list[dict[str, str]] = []
    for row in rows:
        padded = row + [""] * max(0, 8 - len(row))
        cert_valid = normalize_space(padded[5])
        result.append(
            {
                "name": padded[0],
                "model": padded[1],
                "code": padded[2],
                "range": padded[3],
                "uncertainty": padded[4],
                "cert_valid": cert_valid,
                "trace": padded[6] if len(padded) > 6 else "",
            }
        )
    return [item for item in result if item.get("name")]


def _resolve_measurement_rows_for_blueprint(context: dict[str, str]) -> list[dict[str, str]]:
    parsed_rows = _parse_measurement_items_rows(context.get("measurement_items", ""))
    normalized_parsed_rows = _normalize_measurement_rows_for_blueprint(parsed_rows)
    if normalized_parsed_rows:
        return normalized_parsed_rows
    catalog_rows = parse_instrument_catalog_rows_json(context.get("instrument_catalog_rows_json", ""))
    return _build_measurement_rows_from_catalog_for_blueprint(catalog_rows)


def _normalize_measurement_rows_for_blueprint(rows: list[dict[str, str]]) -> list[dict[str, str]]:
    result: list[dict[str, str]] = []
    for row in rows:
        normalized_row = {
            "name": normalize_space(row.get("name", "")),
            "model": normalize_space(row.get("model", "")),
            "code": normalize_space(row.get("code", "")),
            "range": normalize_space(row.get("range", "")),
            "uncertainty": normalize_space(row.get("uncertainty", "")),
            "cert_valid": normalize_space(row.get("cert_valid", "")),
            "trace": normalize_space(row.get("trace", "")),
        }
        if normalized_row.get("name"):
            result.append(normalized_row)
    return result


def _build_measurement_rows_from_catalog_for_blueprint(
    catalog_rows: list[dict[str, str]],
) -> list[dict[str, str]]:
    result: list[dict[str, str]] = []
    for item in catalog_rows:
        cert = normalize_space(item.get("certificate_no", ""))
        valid = normalize_space(item.get("valid_date", ""))
        cert_valid = " ".join([x for x in (cert, valid) if x]).strip()
        row = {
            "name": normalize_space(item.get("name", "")),
            "model": normalize_space(item.get("model", "")),
            "code": normalize_space(item.get("code", "")),
            "range": normalize_space(item.get("measurement_range", "")),
            "uncertainty": normalize_space(item.get("uncertainty", "")),
            "cert_valid": cert_valid,
            "trace": normalize_space(item.get("traceability_institution", "")),
        }
        if row.get("name"):
            result.append(row)
    return result


def _normalize_general_check_lines_for_blueprint(text: str) -> list[str]:
    source = normalize_multiline_text_preserve_tabs(text, normalize_space=normalize_space)
    if not source:
        return []
    lines: list[str] = []
    for raw in source.splitlines():
        line = str(raw or "").strip()
        if not line:
            continue
        if re.search(r"校准结果\s*/\s*说明|Results of calibration and additional explanation", line, flags=re.IGNORECASE):
            continue
        if re.search(r"^注[:：]?", line):
            break
        if line.startswith("序号/标记") or line.startswith("内容"):
            continue
        if "\t" in line:
            parts = [normalize_space(p) for p in line.split("\t")]
            parts = [p for p in parts if p]
            if not parts:
                continue
            line = " ".join(parts)
        line = normalize_space(line)
        if not line:
            continue
        lines.append(line)
    return lines


def _score_detail_general_check_text(text: str) -> int:
    return score_detail_general_check_text(
        text,
        normalize_space=normalize_space,
        extract_source_general_check_lines=extract_source_general_check_lines,
    )


def _is_detail_general_check_sparse(text: str) -> bool:
    return is_detail_general_check_sparse(text, normalize_space=normalize_space)


def _resolve_detail_general_check_for_generic_fill(context: dict[str, str]) -> str:
    return resolve_detail_general_check_for_generic_fill(
        context,
        normalize_space=normalize_space,
        extract_source_general_check_lines=extract_source_general_check_lines,
    )


def build_r801b_payload(
    context: dict[str, str],
    source_file_path: Path | None,
) -> dict[str, Any]:
    raw_text = context.get("raw_record", "") or ""
    tables: list[list[list[str]]] = []
    if source_file_path and source_file_path.exists() and source_file_path.suffix.lower() == ".docx":
        tables = read_docx_tables(source_file_path)
        if not raw_text:
            raw_text = extract_docx_text(source_file_path)

    text_block = raw_text or _tables_to_text_block(tables)
    instrument_catalog_rows = parse_instrument_catalog_rows_json(context.get("instrument_catalog_rows_json", ""))
    instrument_catalog_tokens = parse_instrument_catalog_tokens(context.get("instrument_catalog_names", ""))
    if not instrument_catalog_tokens and instrument_catalog_rows:
        instrument_catalog_tokens = {normalize_catalog_token(row.get("name", "")) for row in instrument_catalog_rows}
        instrument_catalog_tokens = {token for token in instrument_catalog_tokens if token}
    instrument_rows = extract_instrument_rows(tables, instrument_catalog_tokens)
    if instrument_rows and instrument_catalog_rows:
        instrument_rows = merge_instrument_rows_with_catalog(instrument_rows, instrument_catalog_rows)

    certificate_no = sanitize_context_value(context.get("certificate_no", "")) or extract_value_from_tables(
        tables,
        labels=("缆专检号", "Certificate series number"),
    ) or extract_value_by_regex(
        text_block,
        patterns=(
            r"缆专检号[:：]?\s*([A-Za-z0-9\-]+)",
            r"Certificate\s*series\s*number[:：]?\s*([A-Za-z0-9\-]+)",
        ),
    )
    client_name = sanitize_context_value(context.get("client_name", "")) or extract_value_from_tables(
        tables,
        labels=("委托单位", "Client"),
    ) or extract_value_by_regex(
        text_block,
        patterns=(r"委托单位[:：]?\s*([^\n|]+)",),
    )
    receive_date = sanitize_context_date(context.get("receive_date", "")) or extract_date_from_text(
        text_block,
        "收样日期",
    )
    calibration_date = sanitize_context_date(context.get("calibration_date", "")) or extract_date_from_text(
        text_block,
        "校准日期",
    )
    publish_date = sanitize_context_date(context.get("publish_date", "")) or sanitize_context_date(
        context.get("release_date", ""),
    ) or extract_date_from_text(
        text_block,
        "发布日期",
    )
    receive_date, calibration_date, publish_date = _resolve_report_dates(
        receive_date=receive_date,
        calibration_date=calibration_date,
        publish_date=publish_date,
    )

    if not instrument_rows:
        if instrument_catalog_rows:
            fallback_token = normalize_catalog_token(context.get("device_name", ""))
            if fallback_token:
                matched = next(
                    (row for row in instrument_catalog_rows if normalize_catalog_token(row.get("name", "")) == fallback_token),
                    None,
                )
                if matched:
                    instrument_rows = [dict(matched)]
        if not instrument_rows:
            fallback_name = sanitize_instrument_cell(context.get("device_name", ""))
            fallback_model = sanitize_instrument_cell(context.get("device_model", ""))
            fallback_code = sanitize_instrument_cell(context.get("device_code", ""))
            fallback_token = normalize_catalog_token(fallback_name)
            if instrument_catalog_tokens and fallback_token and fallback_token not in instrument_catalog_tokens:
                fallback_name = ""
            if fallback_name or fallback_model or fallback_code:
                instrument_rows = [
                    {
                        "name": fallback_name,
                        "model": fallback_model,
                        "code": fallback_code,
                        "measurement_range": "",
                        "uncertainty": "",
                        "certificate_no": "",
                        "valid_date": "",
                        "traceability_institution": "",
                    }
                ]
    instrument_rows = instrument_rows[:MAX_INSTRUMENT_ROWS]

    payload = {
        "certificate_no": normalize_space(certificate_no),
        "client_name": normalize_space(client_name),
        "receive_date": receive_date,
        "calibration_date": calibration_date,
        "publish_date": publish_date,
        "manufacturer": sanitize_context_value(context.get("manufacturer", "")),
        "instrument_rows": instrument_rows,
    }
    has_value = any(
        [
            payload["certificate_no"],
            payload["client_name"],
            payload["receive_date"],
            payload["calibration_date"],
            payload["publish_date"],
            payload["manufacturer"],
            payload["instrument_rows"],
        ]
    )
    if not has_value:
        return {}
    return payload


def build_r803b_payload(
    context: dict[str, str],
    source_file_path: Path | None,
) -> dict[str, Any]:
    raw_text = context.get("raw_record", "") or ""
    tables: list[list[list[str]]] = []
    if source_file_path and source_file_path.exists() and source_file_path.suffix.lower() == ".docx":
        tables = read_docx_tables(source_file_path)
        if not raw_text:
            raw_text = extract_docx_text(source_file_path)

    text_block = raw_text or _tables_to_text_block(tables)
    instrument_rows = extract_instrument_rows(tables)
    first_instrument = instrument_rows[0] if instrument_rows else {}

    device_name = sanitize_context_value(context.get("device_name", "")) or extract_value_from_tables(
        tables,
        labels=("器具名称", "设备名称", "仪器名称", "Instrument name"),
    ) or extract_value_by_regex(
        text_block,
        patterns=(r"(?:器具名称|设备名称|仪器名称)[:：]?\s*([^\n|]+)",),
    )
    manufacturer = sanitize_context_value(context.get("manufacturer", "")) or extract_value_from_tables(
        tables,
        labels=("制造厂/商", "制造商", "生产厂商", "厂商", "Manufacturer"),
    ) or extract_value_by_regex(
        text_block,
        patterns=(r"(?:制造厂/商|制造商|生产厂商|厂商|厂家)[:：]?\s*([^\n|]+)",),
    )
    client_name = sanitize_context_value(context.get("client_name", "")) or extract_value_from_tables(
        tables,
        labels=("委托单位", "客户名称", "送校单位", "Client"),
    ) or extract_value_by_regex(
        text_block,
        patterns=(r"(?:委托单位|客户名称|送校单位|Client)[:：]?\s*([^\n|]+)",),
    )
    device_model = sanitize_context_value(context.get("device_model", "")) or extract_value_from_tables(
        tables,
        labels=("型号/规格", "型号规格", "型号", "规格型号", "Model/Specification"),
    ) or extract_value_by_regex(
        text_block,
        patterns=(r"(?:型号/规格|型号规格|规格型号|型号)[:：]?\s*([^\n|]+)",),
    )
    device_code = sanitize_context_value(context.get("device_code", "")) or extract_value_from_tables(
        tables,
        labels=("器具编号", "设备编号", "仪器编号", "出厂编号", "Instrument serial number"),
    ) or extract_value_by_regex(
        text_block,
        patterns=(r"(?:器具编号|设备编号|仪器编号|出厂编号)[:：]?\s*([^\n|]+)",),
    )

    model_code_combined = extract_value_from_tables(
        tables,
        labels=("型号/编号", "型号编号", "型号/器具编号", "Model/Number"),
    ) or extract_value_by_regex(
        text_block,
        patterns=(r"(?:型号/编号|型号编号|型号/器具编号|Model/Number)[:：]?\s*([^\n|]+)",),
    )
    combo_model, combo_code = _split_model_code_combined(model_code_combined)
    if not device_model:
        device_model = combo_model
    if not device_code:
        device_code = combo_code
    if (
        combo_model
        and combo_code
        and normalize_space(device_model) == normalize_space(model_code_combined)
    ):
        device_model = combo_model
        if not device_code:
            device_code = combo_code

    if device_model and not device_code:
        inferred_model, inferred_code = _split_model_code_combined(device_model)
        if inferred_code:
            device_model = inferred_model
            device_code = inferred_code

    if device_code and not device_model:
        inferred_model, inferred_code = _split_model_code_combined(device_code)
        if inferred_model:
            device_model = inferred_model
            device_code = inferred_code

    if normalize_catalog_token(device_name) in {"", "name", "instrumentname", "devicename"}:
        device_name = ""
    if normalize_catalog_token(device_model) in {"", "model", "modelspecification", "modelnumber", "number"}:
        device_model = ""
    if normalize_catalog_token(device_code) in {"", "code", "number", "serialnumber"}:
        device_code = ""
    if (not device_name or _looks_like_label(device_name)) and first_instrument.get("name"):
        device_name = first_instrument.get("name", "")
    if (not device_model or _looks_like_label(device_model)) and first_instrument.get("model"):
        device_model = first_instrument.get("model", "")
    if (not device_code or _looks_like_label(device_code)) and first_instrument.get("code"):
        device_code = first_instrument.get("code", "")
    device_name = sanitize_context_value(device_name)
    manufacturer = sanitize_context_value(manufacturer)
    device_model = sanitize_context_value(device_model)
    device_code = sanitize_context_value(device_code)

    certificate_no = sanitize_context_value(context.get("certificate_no", "")) or extract_value_from_tables(
        tables,
        labels=("缆专检号", "证书编号", "证书号", "Certificate series number"),
    ) or extract_value_by_regex(
        text_block,
        patterns=(
            r"缆专检号[:：]?\s*([A-Za-z0-9\-]+)",
            r"证书(?:编号|号)[:：]?\s*([A-Za-z0-9\-]+)",
            r"Certificate\s*series\s*number[:：]?\s*([A-Za-z0-9\-]+)",
        ),
    )

    address = sanitize_context_value(context.get("address", "")) or extract_value_by_regex(
        text_block,
        patterns=(r"地址[:：]?\s*([^\n|]+)",),
    )
    location = sanitize_context_value(context.get("location", "")) or _extract_location_from_other_calibration_info(
        text_block,
    ) or extract_value_from_tables(
        tables,
        labels=("校准地点", "检测地点", "地点", "Location"),
    ) or extract_value_by_regex(
        text_block,
        patterns=(r"(?:校准地点|检测地点|地点)[:：]?\s*([^\n|]+?)(?:Location|$)",),
    ) or address
    location = _sanitize_location_text(location)
    temperature = normalize_space(context.get("temperature", "")) or _extract_temperature_from_other_calibration_info(
        text_block,
    ) or extract_value_from_tables(
        tables,
        labels=("温度", "Temperature"),
    ) or extract_value_by_regex(
        text_block,
        patterns=(r"温度[:：]?\s*([+-]?[0-9]+(?:\.[0-9]+)?)\s*(?:℃|°C|C)?",),
    )
    humidity = normalize_space(context.get("humidity", "")) or _extract_humidity_from_other_calibration_info(
        text_block,
    ) or extract_value_from_tables(
        tables,
        labels=("湿度", "Humidity"),
    ) or extract_value_by_regex(
        text_block,
        patterns=(r"湿度[:：]?\s*([0-9]+(?:\.[0-9]+)?)\s*%?\s*(?:RH|rh)?",),
    )

    section2_u = normalize_space(context.get("section2_u_mm", "")) or _extract_section_uncertainty(
        text_block,
        "中间铁块直径",
        "mm",
    )
    section2_value = normalize_space(context.get("section2_value_mm", "")) or _extract_section_measured_value(
        text_block,
        "中间铁块直径",
        "mm",
    )
    section3_u = normalize_space(context.get("section3_u_g", "")) or _extract_section_uncertainty(
        text_block,
        "中间铁块质量",
        "g",
    )
    section3_value = normalize_space(context.get("section3_value_g", "")) or _extract_section_measured_value(
        text_block,
        "中间铁块质量",
        "g",
    )
    section4_u = normalize_space(context.get("section4_u_g", "")) or _extract_section_uncertainty(
        text_block,
        "铁锤质量",
        "g",
    )

    basis_items_from_context = _extract_standard_codes_from_context_items(context.get("basis_standard_items", ""))
    basis_standard = sanitize_context_value(context.get("basis_standard", "")) or sanitize_context_value(
        context.get("calibration_basis", ""),
    ) or extract_value_by_regex(
        text_block,
        patterns=(
            r"([A-Za-z]{1,5}\s*/\s*T\s*\d+(?:\.\d+)?-\d{4}[^\n|]*)",
            r"本次校准所依据的技术规范(?:[（(]代号、名称[）)])?[:：]?\s*([^\n|]+)",
            r"(?:检测|校准)依据[:：]?\s*([^\n|]+)",
        ),
    )
    basis_codes = basis_items_from_context or _extract_standard_codes(basis_standard)
    if basis_codes:
        basis_standard = "、".join(basis_codes)
    basis_mode = _normalize_basis_mode(context.get("basis_mode", ""))
    if not basis_mode:
        basis_mode = _infer_basis_mode(text_block)

    hammer_actual_rows = extract_hammer_actual_rows(tables)
    if not hammer_actual_rows:
        hammer_actual_rows = extract_hammer_actual_rows_from_text(text_block)
    context_hammer_rows = extract_hammer_actual_rows_from_context(context)
    if context_hammer_rows:
        hammer_actual_rows = merge_hammer_actual_rows(hammer_actual_rows, context_hammer_rows)

    payload = {
        "device_name": normalize_space(device_name),
        "client_name": normalize_space(client_name),
        "address": normalize_space(address),
        "manufacturer": normalize_space(manufacturer),
        "device_model": normalize_space(device_model),
        "device_code": normalize_space(device_code),
        "certificate_no": normalize_space(certificate_no),
        "basis_standard": normalize_space(basis_standard),
        "basis_mode": basis_mode,
        "location": normalize_space(location),
        "temperature": normalize_space(temperature),
        "humidity": normalize_space(humidity),
        "section2_u_mm": normalize_space(section2_u),
        "section2_value_mm": normalize_space(section2_value),
        "section3_u_g": normalize_space(section3_u),
        "section3_value_g": normalize_space(section3_value),
        "section4_u_g": normalize_space(section4_u),
        "hammer_actual_rows": hammer_actual_rows,
    }
    has_value = any(
        [
            payload["device_name"],
            payload["client_name"],
            payload["address"],
            payload["manufacturer"],
            payload["device_model"],
            payload["device_code"],
            payload["certificate_no"],
            payload["basis_standard"],
            payload["basis_mode"],
            payload["location"],
            payload["temperature"],
            payload["humidity"],
            payload["section2_u_mm"],
            payload["section2_value_mm"],
            payload["section3_u_g"],
            payload["section3_value_g"],
            payload["section4_u_g"],
            payload["hammer_actual_rows"],
        ]
    )
    if not has_value:
        return {}
    return payload


def build_r825b_payload(
    context: dict[str, str],
    source_file_path: Path | None,
) -> dict[str, Any]:
    payload = build_r803b_payload(context=context, source_file_path=source_file_path)
    receive_date = sanitize_context_date(context.get("receive_date", ""))
    calibration_date = sanitize_context_date(context.get("calibration_date", ""))
    publish_date = sanitize_context_date(context.get("publish_date", "")) or sanitize_context_date(
        context.get("release_date", ""),
    )
    result = {
        "device_name": normalize_space(payload.get("device_name", "")),
        "client_name": normalize_space(payload.get("client_name", "")),
        "address": normalize_space(payload.get("address", "")),
        "manufacturer": normalize_space(payload.get("manufacturer", "")),
        "device_model": normalize_space(payload.get("device_model", "")),
        "device_code": normalize_space(payload.get("device_code", "")),
        "certificate_no": normalize_space(payload.get("certificate_no", "")),
        "basis_standard": normalize_space(payload.get("basis_standard", "")),
        "basis_mode": normalize_space(payload.get("basis_mode", "")),
        "location": normalize_space(payload.get("location", "")),
        "temperature": normalize_space(payload.get("temperature", "")),
        "humidity": normalize_space(payload.get("humidity", "")),
        "receive_date": receive_date,
        "calibration_date": calibration_date,
        "publish_date": publish_date,
    }
    if any(result.values()):
        return result
    return _build_r825b_payload_from_raw_record_fallback(context)


def _build_r825b_payload_from_raw_record_fallback(context: dict[str, str]) -> dict[str, Any]:
    raw_text = str(context.get("raw_record", "") or "").strip()
    if not raw_text:
        return {}
    extracted = extract_fields(raw_text)
    payload = {
        "device_name": normalize_space(extracted.get("device_name", "")),
        "client_name": normalize_space(extracted.get("client_name", "")),
        "address": normalize_space(extracted.get("address", "")),
        "manufacturer": normalize_space(extracted.get("manufacturer", "")),
        "device_model": normalize_space(extracted.get("device_model", "")),
        "device_code": normalize_space(extracted.get("device_code", "")),
        "certificate_no": normalize_space(extracted.get("certificate_no", "")),
        "basis_standard": "",
        "basis_mode": normalize_space(extracted.get("basis_mode", "")),
        "location": normalize_space(extracted.get("location", "")),
        "temperature": normalize_space(extracted.get("temperature", "")),
        "humidity": normalize_space(extracted.get("humidity", "")),
        "receive_date": sanitize_context_date(extracted.get("receive_date", "")),
        "calibration_date": sanitize_context_date(extracted.get("calibration_date", "")),
        "publish_date": sanitize_context_date(extracted.get("release_date", "")),
    }
    if any(payload.values()):
        return payload
    return {}
def build_r802b_payload(
    context: dict[str, str],
    source_file_path: Path | None,
) -> dict[str, Any]:
    base_payload = build_r825b_payload(context=context, source_file_path=source_file_path)
    raw_text = context.get("raw_record", "") or ""
    if not raw_text and source_file_path and source_file_path.exists() and source_file_path.suffix.lower() == ".docx":
        raw_text = extract_docx_text(source_file_path)

    detail_instruments = extract_text_block(
        raw_text,
        start_patterns=(r"本次校准所使用的主要计量标准器具", r"主要计量标准器具", r"Main measurement standard instruments"),
        end_patterns=(r"本次校准所依据的技术规范", r"(?:其它|其他)校准信息", r"一般检查", r"备注"),
        normalize_space=normalize_space,
    )
    detail_basis = _extract_basis_detail_text(raw_text)
    detail_calibration_info = extract_text_block(
        raw_text,
        start_patterns=(r"(?:其它|其他)校准信息", r"Calibration Information"),
        end_patterns=(r"一般检查", r"备注", r"结果", r"检测员", r"校准员", r"核验员"),
        normalize_space=normalize_space,
    )
    detail_general_check = normalize_multiline_text_preserve_tabs(context.get("general_check_full", ""), normalize_space=normalize_space) or normalize_multiline_text(context.get("general_check", ""), normalize_space=normalize_space) or extract_text_block(
        raw_text,
        start_patterns=(r"(?:一[、.．)]\s*)?一般检查", r"General inspection"),
        end_patterns=(r"^\s*(?:二|2)[、.．)]", r"备注", r"结果", r"检测员", r"校准员", r"核验员"),
        normalize_space=normalize_space,
    )

    payload = {
        **base_payload,
        "detail_instruments": detail_instruments,
        "detail_basis": detail_basis,
        "detail_calibration_info": detail_calibration_info,
        "detail_general_check": detail_general_check,
    }
    if any(payload.values()):
        return payload
    return {}


def _extract_basis_detail_text(text: str) -> str:
    basis_text = extract_text_block(
        text,
        start_patterns=(
            r"本次校准所依据的技术规范",
            r"本次校准所依据的技术规范[（(]代号、名称[）)]",
            r"Reference documents for the calibration",
            r"(?:检测|校准)依据",
        ),
        end_patterns=(
            r"(?:其它|其他)校准信息",
            r"Calibration Information",
            r"(?:一[、.．)]\s*)?一般检查",
            r"General inspection",
            r"备注",
            r"结果",
            r"检测员",
            r"校准员",
            r"核验员",
        ),
        normalize_space=normalize_space,
    )
    if not basis_text:
        return ""
    basis_codes = _extract_standard_codes(basis_text)
    if basis_codes:
        return "\n".join(basis_codes)
    return basis_text


def _find_r802b_record_table(tables: list[ET.Element]) -> ET.Element | None:
    for tbl in tables:
        row_text = " ".join([get_cell_text(tc) for tc in tbl.findall("./w:tr/w:tc", NS)])
        if "原 始 记 录" in row_text and "校准依据" in row_text and "器具编号" in row_text:
            return tbl
    return None


def _fill_r802b_record_table(tbl: ET.Element, payload: dict[str, Any]) -> None:
    rows = tbl.findall("./w:tr", NS)
    if not rows:
        return

    basis_mode = _normalize_basis_mode(payload.get("basis_mode", ""))
    for tr in rows:
        cells = tr.findall("./w:tc", NS)
        if not cells:
            continue

        if payload.get("certificate_no"):
            serial_idx = _find_cell_index_contains(cells, "序")
            if serial_idx >= 0 and "号" in get_cell_text(cells[serial_idx]):
                set_cell_text(cells[serial_idx], f"序 号：{payload['certificate_no']}")

        if payload.get("basis_standard") or basis_mode:
            basis_idx = _find_cell_index_contains(cells, "校准依据")
            if basis_idx >= 0:
                current = get_cell_text(cells[basis_idx])
                basis_text = normalize_space(payload.get("basis_standard", "")) or _extract_basis_from_cell(current)
                set_cell_text(cells[basis_idx], f"{_format_dual_mode_checkbox(basis_mode)}依据：{basis_text}")

        if payload.get("device_name"):
            name_idx = _find_cell_index_contains(cells, "器具名称")
            if name_idx >= 0:
                set_cell_text(cells[name_idx], f"器具名称：{payload['device_name']}")

        if payload.get("manufacturer"):
            manufacturer_idx = _find_cell_index_contains(cells, "制造厂/商")
            if manufacturer_idx >= 0:
                set_cell_text(cells[manufacturer_idx], f"制造厂/商：{payload['manufacturer']}")

        if payload.get("device_model"):
            model_idx = _find_cell_index_contains(cells, "型号/规格")
            if model_idx >= 0:
                set_cell_text(cells[model_idx], f"型号/规格：{payload['device_model']}")

        if payload.get("device_code"):
            code_idx = _find_cell_index_contains(cells, "器具编号")
            if code_idx >= 0:
                set_cell_text(cells[code_idx], f"器具编号：{payload['device_code']}")

        location_idx = _find_cell_index_contains(cells, "校准地点")
        if location_idx >= 0 and (basis_mode or payload.get("location")):
            set_cell_text(
                cells[location_idx],
                f"{_format_dual_mode_checkbox(basis_mode)}地点：{payload.get('location', '')}",
            )

        if payload.get("temperature"):
            _fill_value_between_markers(
                cells=cells,
                start_marker="温度",
                end_marker="℃",
                value=payload["temperature"],
            )

        if payload.get("humidity"):
            _fill_value_between_markers(
                cells=cells,
                start_marker="湿度",
                end_marker="%RH",
                value=payload["humidity"],
            )


def _append_r802b_detail_blocks(
    root: ET.Element,
    payload: dict[str, Any],
    include_general_check: bool = True,
) -> None:
    body = root.find("./w:body", NS)
    if body is None:
        return

    sections: list[tuple[str, str]] = [
        ("本次校准所使用的主要计量标准器具：", normalize_multiline_text(payload.get("detail_instruments", ""), normalize_space=normalize_space)),
        ("本次校准所依据的技术规范（代号、名称）：", normalize_multiline_text(payload.get("detail_basis", ""), normalize_space=normalize_space)),
        ("其它校准信息：", normalize_multiline_text(payload.get("detail_calibration_info", ""), normalize_space=normalize_space)),
    ]
    if include_general_check:
        sections.append(("一般检查：", normalize_multiline_text(payload.get("detail_general_check", ""), normalize_space=normalize_space)))

    section_lines: list[str] = []
    for title, value in sections:
        if not value:
            continue
        if section_lines:
            section_lines.append("")
        if _has_r802b_section_heading(value, title):
            section_lines.extend(value.splitlines())
            continue
        section_lines.append(title)
        section_lines.extend(value.splitlines())
    if not section_lines:
        return

    insert_index = _find_r802b_insert_index(body)
    for line in section_lines:
        paragraph = ET.Element(f"{{{W_NS}}}p")
        if line:
            run = ET.SubElement(paragraph, f"{{{W_NS}}}r")
            text = ET.SubElement(run, f"{{{W_NS}}}t")
            text.text = line
        body.insert(insert_index, paragraph)
        insert_index += 1


def _copy_r802b_general_check_table_from_source(
    target_root: ET.Element,
    source_file_path: Path | None,
) -> tuple[bool, list[ET.Element]]:
    if source_file_path is None or not source_file_path.exists() or source_file_path.suffix.lower() != ".docx":
        return False, []
    try:
        with zipfile.ZipFile(source_file_path, "r") as zf:
            source_xml = zf.read(DOC_XML_PATH)
        source_root = ET.fromstring(source_xml)
    except Exception:
        return False, []

    source_body = source_root.find("./w:body", NS)
    if source_body is None:
        return False, []

    source_range = _find_r802b_general_check_body_range(source_body)
    if source_range is None:
        return False, []
    source_start, source_end = source_range
    source_children = list(source_body)
    source_block = source_children[source_start:source_end]
    if not source_block:
        return False, []

    body = target_root.find("./w:body", NS)
    if body is None:
        return False, []

    insert_index = _find_r802b_general_check_insert_index(body)
    copied_nodes: list[ET.Element] = []
    for node in source_block:
        cloned_node = ET.fromstring(ET.tostring(node, encoding="utf-8"))
        _strip_general_check_required_marker(cloned_node)
        # Keep source continued page block as-is for faithful rendering.
        if _should_skip_r802b_copied_node(cloned_node):
            continue
        _trim_r802b_note_rows_in_node(cloned_node)
        if cloned_node.tag == f"{{{W_NS}}}tbl":
            _sanitize_general_check_table_rows(cloned_node)
        if cloned_node.tag == f"{{{W_NS}}}tbl" and _is_empty_table_after_sanitize(cloned_node):
            continue
        body.insert(insert_index, cloned_node)
        copied_nodes.append(cloned_node)
        insert_index += 1
    if not copied_nodes:
        return False, []
    return True, copied_nodes


def _find_r802b_general_check_body_range(body: ET.Element) -> tuple[int, int] | None:
    children = list(body)
    if not children:
        return None

    end = len(children)
    for idx, child in enumerate(children):
        if child.tag == f"{{{W_NS}}}sectPr":
            end = idx
            break

    start = -1
    for idx in range(end):
        text = normalize_space(" ".join([(t.text or "") for t in children[idx].findall(".//w:t", NS)]))
        if not text:
            continue
        if re.search(r"一般检查|General inspection", text, flags=re.IGNORECASE):
            start = idx
            break
        if re.search(r"校准结果\s*/\s*说明|Results of calibration and additional explanation", text, flags=re.IGNORECASE):
            start = idx
            break
    if start < 0 or start >= end:
        return None
    return start, end


def _copy_modify_certificate_continued_page_table_from_source(
    target_root: ET.Element,
    source_file_path: Path | None,
) -> tuple[bool, list[ET.Element]]:
    if source_file_path is None or not source_file_path.exists() or source_file_path.suffix.lower() != ".docx":
        return False, []
    try:
        with zipfile.ZipFile(source_file_path, "r") as zf:
            source_xml = zf.read(DOC_XML_PATH)
        source_root = ET.fromstring(source_xml)
    except Exception:
        return False, []

    source_body = source_root.find("./w:body", NS)
    target_body = target_root.find("./w:body", NS)
    if source_body is None or target_body is None:
        return False, []

    source_range = _find_modify_certificate_continued_body_range(source_body)
    target_range = _find_modify_certificate_continued_body_range(target_body)
    if source_range is None or target_range is None:
        return False, []

    source_start, source_end = source_range
    target_start, target_end = target_range

    source_children = list(source_body)
    target_children = list(target_body)
    source_block = source_children[source_start:source_end]
    if not source_block:
        return False, []

    for idx in range(target_end - 1, target_start - 1, -1):
        target_body.remove(target_children[idx])

    copied_nodes: list[ET.Element] = []
    insert_idx = target_start
    for node in source_block:
        cloned_node = ET.fromstring(ET.tostring(node, encoding="utf-8"))
        target_body.insert(insert_idx, cloned_node)
        copied_nodes.append(cloned_node)
        insert_idx += 1

    copied_tables: list[ET.Element] = []
    for node in copied_nodes:
        if node.tag == f"{{{W_NS}}}tbl":
            copied_tables.append(node)
        copied_tables.extend(node.findall(".//w:tbl", NS))

    if not copied_nodes:
        return False, []
    return True, copied_tables


def _find_modify_certificate_continued_body_range(body: ET.Element) -> tuple[int, int] | None:
    children = list(body)
    if not children:
        return None

    end = len(children)
    for idx, child in enumerate(children):
        if child.tag == f"{{{W_NS}}}sectPr":
            end = idx
            break

    start = -1
    for idx in range(end):
        if _is_modify_certificate_continued_heading_child(children[idx]):
            start = idx
            break
    if start < 0 or start >= end:
        return None
    return start, end


def _is_modify_certificate_continued_heading_child(node: ET.Element) -> bool:
    text = normalize_space(" ".join([(t.text or "") for t in node.findall(".//w:t", NS)]))
    if not text:
        return False
    return bool(re.search(r"校准结果\s*/\s*说明|Results of calibration and additional explanation", text, flags=re.IGNORECASE))


def _find_modify_certificate_continued_page_table(root: ET.Element) -> ET.Element | None:
    tables = _find_modify_certificate_continued_page_tables(root)
    if tables:
        return tables[0]

    candidates: list[tuple[int, ET.Element]] = []
    for tbl in root.findall(".//w:tbl", NS):
        score = _score_modify_certificate_continued_page_table(tbl)
        if score < -9999:
            continue
        candidates.append((score, tbl))
    if not candidates:
        return None
    candidates.sort(key=lambda item: item[0], reverse=True)
    return candidates[0][1]


def _find_modify_certificate_continued_page_tables(root: ET.Element) -> list[ET.Element]:
    selected: list[ET.Element] = []
    candidates: list[tuple[int, ET.Element]] = []
    for tbl in root.findall(".//w:tbl", NS):
        score = _score_modify_certificate_continued_page_table(tbl)
        if score < -9999:
            continue
        candidates.append((score, tbl))
        if _is_modify_certificate_continued_page_table_candidate(tbl):
            selected.append(tbl)
    if selected:
        return selected
    if not candidates:
        return []
    candidates.sort(key=lambda item: item[0], reverse=True)
    return [candidates[0][1]]


def _is_modify_certificate_continued_page_table_candidate(tbl: ET.Element) -> bool:
    rows = tbl.findall("./w:tr", NS)
    row_count = len(rows)
    if row_count < 2:
        return False
    text = normalize_space(" ".join([(node.text or "") for node in tbl.findall(".//w:t", NS)]))
    if not text:
        return False
    if re.search(r"一般检查|General inspection", text, flags=re.IGNORECASE):
        return True
    if re.search(r"\(\s*1\s*\)|（\s*1\s*）", text):
        return True
    if re.search(r"\b注[:：]?|\bNotes?[:：]?", text, flags=re.IGNORECASE):
        return True
    max_cols = 0
    for row in rows:
        cols = len(row.findall("./w:tc", NS))
        if cols > max_cols:
            max_cols = cols
    return max_cols >= 2


def _score_modify_certificate_continued_page_table(tbl: ET.Element) -> int:
    text = normalize_space(" ".join([(node.text or "") for node in tbl.findall(".//w:t", NS)]))
    if not text:
        return -10000
    if not re.search(r"校准结果\s*/\s*说明|Results of calibration and additional explanation", text, flags=re.IGNORECASE):
        return -10000
    rows = tbl.findall("./w:tr", NS)
    row_count = len(rows)
    max_cols = 0
    for row in rows:
        cols = len(row.findall("./w:tc", NS))
        if cols > max_cols:
            max_cols = cols
    score = row_count * 5
    if re.search(r"一般检查|General inspection", text, flags=re.IGNORECASE):
        score += 80
    if re.search(r"\(\s*1\s*\)|（\s*1\s*）", text):
        score += 60
    if re.search(r"^注[:：]?|^Notes?[:：]?", text, flags=re.IGNORECASE):
        score += 60
    if max_cols >= 3:
        score += 50
    if row_count <= 2:
        score -= 120
    return score


def _insert_xml_element_after(root: ET.Element, anchor_elem: ET.Element, new_elem: ET.Element) -> bool:
    for parent in root.iter():
        children = list(parent)
        for idx, child in enumerate(children):
            if child is anchor_elem:
                parent.insert(idx + 1, new_elem)
                return True
    return False


def _replace_xml_element(root: ET.Element, old_elem: ET.Element, new_elem: ET.Element) -> bool:
    for parent in root.iter():
        children = list(parent)
        for idx, child in enumerate(children):
            if child is old_elem:
                parent.remove(old_elem)
                parent.insert(idx, new_elem)
                return True
    return False


def _find_general_check_table_element(root: ET.Element) -> ET.Element | None:
    candidates: list[tuple[int, ET.Element]] = []
    for tbl in root.findall(".//w:tbl", NS):
        text = normalize_space(" ".join([(node.text or "") for node in tbl.findall(".//w:t", NS)]))
        if not text:
            continue
        if not re.search(r"一般检查|General inspection", text, flags=re.IGNORECASE):
            continue
        if not re.search(r"显示值|实测值|试验温度校准|注[:：]|\(\s*1\s*\)|（\s*1\s*）", text):
            continue

        score = 0
        if re.search(r"校准结果\s*/\s*说明|Results of calibration and additional explanation", text, flags=re.IGNORECASE):
            score -= 100
        if re.search(r"本次校准所使用的主要计量标准器具|本次校准所依据的技术规范|(?:其它|其他)校准信息", text):
            score -= 100
        score += len(tbl.findall(".//w:tr", NS))
        candidates.append((score, tbl))

    if not candidates:
        return None
    candidates.sort(key=lambda item: item[0], reverse=True)
    return candidates[0][1]


def _sanitize_general_check_table_rows(tbl: ET.Element) -> None:
    rows = list(tbl.findall("./w:tr", NS))
    if not rows:
        return

    start_index = -1
    for idx, row in enumerate(rows):
        text = normalize_space(" ".join([(node.text or "") for node in row.findall(".//w:t", NS)]))
        if re.search(r"一般检查|General inspection", text, flags=re.IGNORECASE):
            start_index = idx
            break
        if re.search(r"(?:^|[\s])(?:[一二三四五六七八九十]+[、.．)]|\(\s*\d+\s*\)|（\s*\d+\s*）)", text):
            start_index = idx
            break

    keep_until = len(rows)
    for idx, row in enumerate(rows):
        text = normalize_space(" ".join([(node.text or "") for node in row.findall(".//w:t", NS)]))
        if re.search(r"注[:：]?|Notes?[:：]?|备注|Remarks|检测员|校准员|核验员|(?:以下空白|\(以下空白\)|（以下空白）)", text, flags=re.IGNORECASE):
            keep_until = idx
            break

    for idx, row in enumerate(rows):
        text = normalize_space(" ".join([(node.text or "") for node in row.findall(".//w:t", NS)]))
        remove = False
        if start_index >= 0 and idx < start_index:
            remove = True
        if idx >= keep_until:
            remove = True
        if re.search(r"校准结果\s*/\s*说明|Results of calibration and additional explanation", text, flags=re.IGNORECASE):
            remove = True
        if re.search(
            r"校准证书续页专用|Continued page of calibration certificate|"
            r"第\s*\d+\s*页\s*[\/／]\s*共\s*\d+\s*页|"
            r"\bPage\b(?:\s+\d+\s+of\s+\d+)?|\bof\s*total\b|"
            r"缆专检号|Certificate series number|"
            r"上海国缆检测股份有限公司|Shanghai National Center of Testing and Inspection",
            text,
            flags=re.IGNORECASE,
        ):
            remove = True
        if re.search(r"本次校准所使用的主要计量标准器具|本次校准所依据的技术规范|(?:其它|其他)校准信息", text):
            remove = True
        if remove:
            tbl.remove(row)


def _should_skip_r802b_continued_node(node: ET.Element) -> bool:
    text = normalize_space(" ".join([(t.text or "") for t in node.findall(".//w:t", NS)]))
    if not text:
        return False
    has_body_marker = bool(
        re.search(
            r"一般检查|General inspection|"
            r"(?:^|[\s])(?:[一二三四五六七八九十]+[、.．)]|\(\s*\d+\s*\)|（\s*\d+\s*）)|"
            r"实测值|扩展不确定度|试验|校准值|显示值",
            text,
            flags=re.IGNORECASE,
        )
    )
    if has_body_marker:
        return False
    return bool(
        re.search(
            r"校准证书续页专用|Continued page of calibration certificate|"
            r"第\s*\d+\s*页\s*[\/／]\s*共\s*\d+\s*页|"
            r"\bPage\b(?:\s+\d+\s+of\s+\d+)?|\bof\s*total\b|"
            r"校准结果\s*/\s*说明\s*[（(]?\s*续页\s*[）)]?\s*[:：]?|"
            r"Results of calibration and additional explanation\s*[（(]?\s*continued page\s*[）)]?\s*[:：]?|"
            r"缆专检号|Certificate series number",
            text,
            flags=re.IGNORECASE,
        )
    )


def _should_skip_r802b_copied_node(node: ET.Element) -> bool:
    text = normalize_space(" ".join([(t.text or "") for t in node.findall(".//w:t", NS)]))
    if not text:
        return False
    # Skip copied page header/footer blocks.
    if re.search(r"校准证书续页专用|Continued page of calibration certificate", text, flags=re.IGNORECASE):
        return True
    if re.search(r"上海国缆检测股份有限公司|Shanghai National Center of Testing and Inspection", text, flags=re.IGNORECASE) and re.search(
        r"缆专检号|Certificate series number",
        text,
        flags=re.IGNORECASE,
    ):
        return True
    # Skip standalone continued-page title block; keep only the content below it.
    if re.search(r"校准结果\s*/\s*说明|Results of calibration and additional explanation", text, flags=re.IGNORECASE):
        has_body_marker = bool(
            re.search(
                r"(?:^|[\s])(?:[一二三四五六七八九十]+[、.．)]|\(\s*\d+\s*\)|（\s*\d+\s*）)|双脉冲|频带宽度|放电量表|衰减系数",
                text,
                flags=re.IGNORECASE,
            )
        )
        if not has_body_marker:
            return True
    return False


def _trim_r802b_note_rows_in_node(node: ET.Element) -> None:
    for tbl in node.findall(".//w:tbl", NS):
        rows = list(tbl.findall("./w:tr", NS))
        if not rows:
            continue
        cut_idx = -1
        for idx, row in enumerate(rows):
            row_text = normalize_space(" ".join([(t.text or "") for t in row.findall(".//w:t", NS)]))
            if not row_text:
                continue
            if re.search(r"注[:：]?|Notes?[:：]?|备注[:：]?|Remarks[:：]?|(?:以下空白|\(以下空白\)|（以下空白）)", row_text, flags=re.IGNORECASE):
                cut_idx = idx
                break
        if cut_idx < 0:
            continue
        for idx in range(len(rows) - 1, cut_idx - 1, -1):
            tbl.remove(rows[idx])


def _is_empty_table_after_sanitize(tbl: ET.Element) -> bool:
    rows = tbl.findall("./w:tr", NS)
    if not rows:
        return True
    for row in rows:
        text = normalize_space(" ".join([(t.text or "") for t in row.findall(".//w:t", NS)]))
        if text:
            return False
    return True


def _append_r802b_general_check_text_only(root: ET.Element, payload: dict[str, Any]) -> None:
    body = root.find("./w:body", NS)
    if body is None:
        return
    value = normalize_multiline_text(payload.get("detail_general_check", ""), normalize_space=normalize_space)
    if not value:
        return

    filtered_lines: list[str] = []
    for line in value.splitlines():
        compact = normalize_space(line)
        if not compact:
            continue
        if re.search(r"校准结果\s*/\s*说明|Results of calibration and additional explanation", compact, flags=re.IGNORECASE):
            continue
        filtered_lines.append(compact)
    if not filtered_lines:
        return

    insert_index = _find_r802b_general_check_insert_index(body)
    for line in filtered_lines:
        paragraph = ET.Element(f"{{{W_NS}}}p")
        run = ET.SubElement(paragraph, f"{{{W_NS}}}r")
        text = ET.SubElement(run, f"{{{W_NS}}}t")
        text.text = line
        body.insert(insert_index, paragraph)
        insert_index += 1


def _find_r802b_general_check_insert_index(body: ET.Element) -> int:
    children = list(body)
    for idx, child in enumerate(children):
        if child.tag != f"{{{W_NS}}}tbl":
            continue
        text = "".join([(node.text or "") for node in child.findall(".//w:t", NS)])
        if "器具编号" in text and "型号/规格" in text:
            return idx + 1
    return _find_r802b_insert_index(body)


def _find_r802b_insert_index(body: ET.Element) -> int:
    children = list(body)
    for idx, child in enumerate(children):
        if child.tag != f"{{{W_NS}}}tbl":
            continue
        text = "".join([(node.text or "") for node in child.findall(".//w:t", NS)])
        if "检测员" in text and "核验员" in text and "页/共" in text:
            return idx

    for idx, child in enumerate(children):
        if child.tag == f"{{{W_NS}}}sectPr":
            return idx
    return len(children)


def _has_r802b_section_heading(value: str, title: str) -> bool:
    text = normalize_multiline_text(value, normalize_space=normalize_space)
    if not text:
        return False
    if title in text:
        return True
    if "一般检查" in title and re.search(r"(?:^|\n)(?:一[、.．)]\s*)?一般检查[:：]?", text):
        return True
    if "其它校准信息" in title and re.search(r"(?:^|\n)(?:其它|其他)校准信息[:：]?", text):
        return True
    if "技术规范" in title and re.search(r"(?:^|\n)本次校准所依据的技术规范", text):
        return True
    if "主要计量标准器具" in title and re.search(r"(?:^|\n)本次校准所使用的主要计量标准器具", text):
        return True
    return False


def build_r803b_editor_fields(
    context: dict[str, str],
    source_file_path: Path | None,
) -> dict[str, str]:
    payload: dict[str, Any] = {}
    # Editor prefill should prefer source-doc parsing so incomplete OCR/extract
    # values (for example truncated manufacturer) do not override full values.
    if source_file_path and source_file_path.exists() and source_file_path.suffix.lower() == ".docx":
        payload = build_r803b_payload(context={}, source_file_path=source_file_path)
    if not payload:
        payload = build_r803b_payload(
            context={"raw_record": context.get("raw_record", "")},
            source_file_path=source_file_path,
        )
    if not payload:
        payload = build_r803b_payload(context=context, source_file_path=source_file_path)
    if not payload:
        return {}

    rows: list[list[str]] = payload.get("hammer_actual_rows", []) or []
    return {
        "device_name": normalize_space(payload.get("device_name", "")),
        "manufacturer": normalize_space(payload.get("manufacturer", "")),
        "device_model": normalize_space(payload.get("device_model", "")),
        "device_code": normalize_space(payload.get("device_code", "")),
        "certificate_no": normalize_space(payload.get("certificate_no", "")),
        "basis_standard": normalize_space(payload.get("basis_standard", "")),
        "basis_mode": normalize_space(payload.get("basis_mode", "")),
        "location": normalize_space(payload.get("location", "")),
        "temperature": normalize_space(payload.get("temperature", "")),
        "humidity": normalize_space(payload.get("humidity", "")),
        "section2_u_mm": normalize_space(payload.get("section2_u_mm", "")),
        "section2_value_mm": normalize_space(payload.get("section2_value_mm", "")),
        "section3_u_g": normalize_space(payload.get("section3_u_g", "")),
        "section3_value_g": normalize_space(payload.get("section3_value_g", "")),
        "section4_u_g": normalize_space(payload.get("section4_u_g", "")),
        "hammer_actual_row_1": " ".join(rows[0]) if len(rows) > 0 else "",
        "hammer_actual_row_2": " ".join(rows[1]) if len(rows) > 1 else "",
        "hammer_actual_row_3": " ".join(rows[2]) if len(rows) > 2 else "",
    }


def _find_info_table(tables: list[ET.Element]) -> ET.Element | None:
    for tbl in tables:
        if _table_contains(tbl, "缆专检号") and _table_contains(tbl, "委托单位"):
            return tbl
    return None


def _find_record_table(tables: list[ET.Element]) -> ET.Element | None:
    for tbl in tables:
        row_text = " ".join([get_cell_text(tc) for tc in tbl.findall("./w:tr/w:tc", NS)])
        if "器具名称" in row_text and "型号/规格" in row_text and "编号" in row_text:
            return tbl
    return None


def _table_contains(tbl: ET.Element, keyword: str) -> bool:
    return keyword in "".join([(node.text or "") for node in tbl.findall(".//w:t", NS)])


def _fill_info_table(tbl: ET.Element, payload: dict[str, Any]) -> None:
    certificate_no = payload.get("certificate_no", "")
    client_name = payload.get("client_name", "")
    receive_date = payload.get("receive_date", "")
    calibration_date = payload.get("calibration_date", "")
    publish_date = payload.get("publish_date", "")

    for tr in tbl.findall("./w:tr", NS):
        cells = tr.findall("./w:tc", NS)
        if not cells:
            continue
        cell_texts = [get_cell_text(cell) for cell in cells]
        if certificate_no:
            for idx, text in enumerate(cell_texts):
                if "缆专检号" in text:
                    set_cell_text(cells[idx], f"缆专检号：{certificate_no}")
        if client_name:
            for idx, text in enumerate(cell_texts):
                if "委托单位" in text:
                    set_cell_text(cells[idx], f"委托单位：{client_name}")
        if any("收样日期" in text for text in cell_texts):
            _fill_split_date(cells, "收样日期", receive_date)
        if any("检测/校准日期" in text for text in cell_texts):
            _fill_split_date(cells, "检测/校准日期", calibration_date)
        if any("校准日期" in text for text in cell_texts):
            _fill_split_date(cells, "校准日期", calibration_date)
        if any("发布日期" in text for text in cell_texts):
            _fill_split_date(cells, "发布日期", publish_date)


def _fill_split_date(cells: list[ET.Element], label: str, date_text: str) -> None:
    if not date_text:
        return
    date_parts = split_date_parts(date_text)
    if not date_parts:
        return
    year, month, day = date_parts

    label_idx = -1
    for idx, cell in enumerate(cells):
        if _contains_compact_label(get_cell_text(cell), label):
            label_idx = idx
            break
    if label_idx < 0:
        return

    range_indices = range(label_idx + 1, len(cells))
    year_marker = _find_cell_index_with_text(cells, range_indices, "年")
    month_marker = _find_cell_index_with_text(cells, range_indices, "月")
    day_marker = _find_cell_index_with_text(cells, range_indices, "日")
    if year_marker > 0:
        set_cell_text(cells[year_marker - 1], year)
    if month_marker > 0:
        set_cell_text(cells[month_marker - 1], month)
    if day_marker > 0:
        set_cell_text(cells[day_marker - 1], day)


def _fill_record_table(tbl: ET.Element, payload: dict[str, Any]) -> None:
    rows = tbl.findall("./w:tr", NS)
    if len(rows) < 2:
        return

    data_rows: list[tuple[ET.Element, list[ET.Element]]] = []
    for tr in rows[1:]:
        cells = tr.findall("./w:tc", NS)
        if len(cells) < 7:
            break
        first_text = get_cell_text(cells[0])
        if first_text.startswith("注："):
            break
        data_rows.append((tr, cells))

    if not data_rows:
        return

    instrument_rows = payload.get("instrument_rows", []) or []
    normalized_instrument_rows = [
        {
            "name": sanitize_instrument_cell(item.get("name", "")),
            "model": sanitize_instrument_cell(item.get("model", "")),
            "code": sanitize_instrument_cell(item.get("code", "")),
            "measurement_range": sanitize_instrument_cell(item.get("measurement_range", "")),
            "uncertainty": sanitize_instrument_cell(item.get("uncertainty", "")),
            "certificate_no": sanitize_instrument_cell(item.get("certificate_no", "")),
            "valid_date": sanitize_instrument_cell(item.get("valid_date", "")),
            "traceability_institution": sanitize_instrument_cell(item.get("traceability_institution", "")),
        }
        for item in instrument_rows
    ]
    normalized_instrument_rows = [item for item in normalized_instrument_rows if item["name"]][:MAX_INSTRUMENT_ROWS]

    header_cells = rows[0].findall("./w:tc", NS)
    index_name = _find_cell_index_contains_any(header_cells, ("器具名称", "Instrument name"))
    index_model = _find_cell_index_contains_any(header_cells, ("型号/规格", "Model/Specification"))
    index_code = _find_cell_index_contains_any(header_cells, ("编号", "Number"))
    index_range = _find_cell_index_contains_any(header_cells, ("测量范围", "Measurement range"))
    index_uncertainty = _find_cell_index_contains_any(header_cells, ("不确定度", "最大允许误差", "Uncertainty"))
    index_cert_valid = _find_cell_index_contains_any(header_cells, ("证书编号", "有效期限", "Certificate number", "Valid date"))
    index_trace = _find_cell_index_contains_any(header_cells, ("溯源机构", "traceability institution"))
    if index_name < 0:
        index_name = 0
    if index_model < 0:
        index_model = 1
    if index_code < 0:
        index_code = 2

    for idx, (tr, cells) in enumerate(data_rows):
        if idx >= len(normalized_instrument_rows):
            tbl.remove(tr)
            continue
        item = normalized_instrument_rows[idx]
        for cell in cells:
            set_cell_text(cell, "")
        if index_name < len(cells):
            set_cell_text(cells[index_name], f"□{item['name']}")
        if index_model < len(cells):
            set_cell_text(cells[index_model], item["model"])
        if index_code < len(cells):
            set_cell_text(cells[index_code], item["code"])
        if index_range >= 0 and index_range < len(cells):
            set_cell_text(cells[index_range], item["measurement_range"])
        if index_uncertainty >= 0 and index_uncertainty < len(cells):
            set_cell_text(cells[index_uncertainty], item["uncertainty"])
        if index_cert_valid >= 0 and index_cert_valid < len(cells):
            cert_parts = [part for part in (item["certificate_no"], item["valid_date"]) if part]
            set_cell_text(cells[index_cert_valid], "\n".join(cert_parts))
        elif len(cells) > 3:
            set_cell_text(cells[3], item["valid_date"])
        if index_trace >= 0 and index_trace < len(cells):
            set_cell_text(cells[index_trace], item["traceability_institution"])

    manufacturer = normalize_space(payload.get("manufacturer", ""))
    if manufacturer and data_rows and len(normalized_instrument_rows) > 0:
        target_cells = data_rows[0][1]
        if index_trace >= 0 and index_trace < len(target_cells) and not get_cell_text(target_cells[index_trace]):
            set_cell_text(target_cells[index_trace], manufacturer)
        elif len(target_cells) > 6:
            set_cell_text(target_cells[6], manufacturer)


def _find_r803b_record_table(tables: list[ET.Element]) -> ET.Element | None:
    for tbl in tables:
        row_text = " ".join([get_cell_text(tc) for tc in tbl.findall("./w:tr/w:tc", NS)])
        if "低温冲击试验装置原始记录" in row_text and "环境条件" in row_text and "铁锤质量" in row_text:
            return tbl
    return None


def _find_r825b_record_table(tables: list[ET.Element]) -> ET.Element | None:
    for tbl in tables:
        row_text = " ".join([get_cell_text(tc) for tc in tbl.findall("./w:tr/w:tc", NS)])
        if "软化击穿试验仪原始记录" in row_text and "环境条件" in row_text and "试样间短路电流" in row_text:
            return tbl
    return None


def _find_generic_record_table(tables: list[ET.Element]) -> ET.Element | None:
    return find_generic_record_table_by_rules(tables=tables, get_cell_text=get_cell_text)


def _fill_r803b_record_table(tbl: ET.Element, payload: dict[str, Any]) -> None:
    rows = tbl.findall("./w:tr", NS)
    if not rows:
        return

    basis_mode = _normalize_basis_mode(payload.get("basis_mode", ""))

    for tr in rows:
        cells = tr.findall("./w:tc", NS)
        if not cells:
            continue
        cell_texts = [get_cell_text(cell) for cell in cells]

        if payload.get("certificate_no"):
            serial_idx = _find_cell_index_contains(cells, "序")
            if serial_idx >= 0 and "号" in get_cell_text(cells[serial_idx]):
                set_cell_text(cells[serial_idx], f"序 号：{payload['certificate_no']}")

        if payload.get("basis_standard") or basis_mode:
            basis_idx = _find_cell_index_contains(cells, "校准依据")
            if basis_idx >= 0:
                current = get_cell_text(cells[basis_idx])
                basis_text = normalize_space(payload.get("basis_standard", "")) or _extract_basis_from_cell(current)
                set_cell_text(cells[basis_idx], f"{_format_dual_mode_checkbox(basis_mode)}依据：{basis_text}")

        if payload.get("device_name"):
            name_idx = _find_cell_index_contains(cells, "器具名称")
            if name_idx >= 0:
                set_cell_text(cells[name_idx], f"器具名称：{payload['device_name']}")

        if payload.get("manufacturer"):
            manufacturer_idx = _find_cell_index_contains(cells, "制造厂/商")
            if manufacturer_idx >= 0:
                set_cell_text(cells[manufacturer_idx], f"制造厂/商：{payload['manufacturer']}")

        if payload.get("device_model"):
            model_idx = _find_cell_index_contains(cells, "型号/规格")
            if model_idx >= 0:
                set_cell_text(cells[model_idx], f"型号/规格：{payload['device_model']}")

        if payload.get("device_code"):
            code_idx = _find_cell_index_contains(cells, "器具编号")
            if code_idx >= 0:
                set_cell_text(cells[code_idx], f"器具编号：{payload['device_code']}")

        location_idx = _find_cell_index_contains(cells, "校准地点")
        if location_idx >= 0 and (basis_mode or payload.get("location")):
            set_cell_text(
                cells[location_idx],
                f"{_format_dual_mode_checkbox(basis_mode)}地点：{payload.get('location', '')}",
            )

        if payload.get("temperature"):
            _fill_value_between_markers(
                cells=cells,
                start_marker="温度",
                end_marker="℃",
                value=payload["temperature"],
            )

        if payload.get("humidity"):
            _fill_value_between_markers(
                cells=cells,
                start_marker="湿度",
                end_marker="%RH",
                value=payload["humidity"],
            )

        if payload.get("section2_u_mm"):
            idx = _find_cell_index_contains(cells, "二、中间铁块直径")
            if idx >= 0:
                source = get_cell_text(cells[idx])
                set_cell_text(cells[idx], _replace_uncertainty_value(source, payload["section2_u_mm"], "mm"))

        if payload.get("section2_value_mm"):
            idx = _find_cell_index_contains(cells, "实测值")
            if idx >= 0 and "mm" in get_cell_text(cells[idx]):
                source = get_cell_text(cells[idx])
                set_cell_text(cells[idx], _replace_measured_value(source, payload["section2_value_mm"], "mm"))

        if payload.get("section3_u_g"):
            idx = _find_cell_index_contains(cells, "三、中间铁块质量")
            if idx >= 0:
                source = get_cell_text(cells[idx])
                set_cell_text(cells[idx], _replace_uncertainty_value(source, payload["section3_u_g"], "g"))

        if payload.get("section3_value_g"):
            idx = _find_cell_index_contains(cells, "实测值")
            if idx >= 0 and "g" in get_cell_text(cells[idx]):
                source = get_cell_text(cells[idx])
                set_cell_text(cells[idx], _replace_measured_value(source, payload["section3_value_g"], "g"))

        if payload.get("section4_u_g"):
            idx = _find_cell_index_contains(cells, "四、铁锤质量")
            if idx >= 0:
                source = get_cell_text(cells[idx])
                set_cell_text(cells[idx], _replace_uncertainty_value(source, payload["section4_u_g"], "g"))

    _fill_r803b_hammer_actual_rows(tbl, payload.get("hammer_actual_rows", []))


def _fill_r825b_record_table(tbl: ET.Element, payload: dict[str, Any]) -> None:
    rows = tbl.findall("./w:tr", NS)
    if not rows:
        return

    basis_mode = _normalize_basis_mode(payload.get("basis_mode", ""))

    for tr in rows:
        cells = tr.findall("./w:tc", NS)
        if not cells:
            continue

        if payload.get("certificate_no"):
            serial_idx = _find_cell_index_contains(cells, "序")
            if serial_idx >= 0 and "号" in get_cell_text(cells[serial_idx]):
                set_cell_text(cells[serial_idx], f"序 号：{payload['certificate_no']}")

        fill_base_fields_in_cells_by_rules(
            cells=cells,
            payload=payload,
            basis_mode=basis_mode,
            get_cell_text=get_cell_text,
            set_cell_text=set_cell_text,
            extract_basis_from_text=_extract_basis_from_cell,
            format_mode_prefix=_format_dual_mode_checkbox,
        )

        if payload.get("temperature"):
            _fill_value_between_markers(
                cells=cells,
                start_marker="温度",
                end_marker="℃",
                value=payload["temperature"],
            )

        if payload.get("humidity"):
            _fill_value_between_markers(
                cells=cells,
                start_marker="湿度",
                end_marker="%RH",
                value=payload["humidity"],
            )


def _fill_generic_base_labels_in_tables(tables: list[ET.Element], payload: dict[str, Any]) -> bool:
    basis_mode = _normalize_basis_mode(payload.get("basis_mode", ""))
    return fill_base_fields_in_tables_by_rules(
        tables=tables,
        payload=payload,
        basis_mode=basis_mode,
        get_cell_text=get_cell_text,
        set_cell_text=set_cell_text,
        extract_basis_from_text=_extract_basis_from_cell,
        format_mode_prefix=_format_dual_mode_checkbox,
    )


def _fill_generic_base_labels_in_paragraphs(root: ET.Element, payload: dict[str, Any]) -> bool:
    basis_mode = _normalize_basis_mode(payload.get("basis_mode", ""))
    return fill_base_fields_in_paragraphs_by_rules(
        root=root,
        payload=payload,
        basis_mode=basis_mode,
        extract_basis_from_text=_extract_basis_from_cell,
        format_mode_prefix=_format_dual_mode_checkbox,
    )


def _write_docx_with_updated_root(
    template_path: Path,
    output_path: Path,
    root: ET.Element,
    original_namespaces: dict[str, str],
    header_payload: dict[str, Any] | None = None,
    rel_updates: dict[str, bytes] | None = None,
) -> None:
    _preserve_original_namespaces(root, original_namespaces)
    updated_xml = ET.tostring(root, encoding="utf-8", xml_declaration=True)

    updates = rel_updates or {}
    updated_rels_xml = updates.get("__rels__")
    updated_ct_xml = updates.get("__ct__")
    with zipfile.ZipFile(template_path, "r") as zin, zipfile.ZipFile(output_path, "w") as zout:
        existing_names = {x.filename for x in zin.infolist()}
        for item in zin.infolist():
            if item.filename == DOC_XML_PATH:
                zout.writestr(item, updated_xml)
            elif _is_header_xml_path(item.filename) and header_payload is not None:
                header_xml = _fill_header_base_fields_xml(zin.read(item.filename), header_payload)
                zout.writestr(item, header_xml)
            elif item.filename == DOC_RELS_PATH and updated_rels_xml is not None:
                zout.writestr(item, updated_rels_xml)
            elif item.filename == CONTENT_TYPES_PATH and updated_ct_xml is not None:
                zout.writestr(item, updated_ct_xml)
            else:
                zout.writestr(item, zin.read(item.filename))

        if updated_rels_xml is not None and DOC_RELS_PATH not in existing_names:
            zout.writestr(DOC_RELS_PATH, updated_rels_xml)
        if updated_ct_xml is not None and CONTENT_TYPES_PATH not in existing_names:
            zout.writestr(CONTENT_TYPES_PATH, updated_ct_xml)
        for path, raw in updates.items():
            if path.startswith("__"):
                continue
            zout.writestr(path, raw)


def _is_header_xml_path(path: str) -> bool:
    return bool(re.match(r"^word/header\d+\.xml$", str(path or "").strip()))


def _fill_text_placeholders_in_root(root: ET.Element, values: dict[str, str]) -> bool:
    changed = False
    valid_values = {k: str(v or "") for k, v in (values or {}).items() if str(k or "").strip()}
    if not valid_values:
        return False
    for node in root.findall(".//w:t", NS):
        text = str(node.text or "")
        if not text:
            continue
        updated = text
        for key, value in valid_values.items():
            if key not in updated:
                continue
            updated = updated.replace(key, value)
        if updated != text:
            node.text = updated
            changed = True
    return changed


def _resolve_yyyymmdd_serial_no(context: dict[str, Any], payload: dict[str, Any]) -> str:
    context_data = context if isinstance(context, dict) else {}
    payload_data = payload if isinstance(payload, dict) else {}
    explicit = normalize_space(
        str(
            context_data.get("report_no")
            or context_data.get("report_number")
            or payload_data.get("report_no")
            or payload_data.get("report_number")
            or "",
        ),
    )
    if explicit:
        return explicit
    record_no = normalize_space(str(context_data.get("record_no") or payload_data.get("record_no") or ""))
    matched = re.match(r"^(\d{8})[-_]?(\d{1,4})$", record_no)
    if matched:
        return f"{matched.group(1)}{str(matched.group(2) or '').zfill(2)}"
    return f"{datetime.now().strftime('%Y%m%d')}01"


def _resolve_iso_date_text(context: dict[str, Any], payload: dict[str, Any]) -> str:
    context_data = context if isinstance(context, dict) else {}
    payload_data = payload if isinstance(payload, dict) else {}
    for key in ("report_date", "company_sign_date", "release_date", "calibration_date", "receive_date"):
        date_text = sanitize_context_date(str(context_data.get(key) or payload_data.get(key) or ""))
        if not date_text:
            continue
        parts = split_date_parts(date_text)
        if not parts:
            continue
        year, month, day = parts
        return f"{year}-{month}-{day}"
    return datetime.now().strftime("%Y-%m-%d")


def _build_jinja_placeholder_values(context: dict[str, Any], payload: dict[str, Any]) -> dict[str, str]:
    context_data = context if isinstance(context, dict) else {}
    payload_data = payload if isinstance(payload, dict) else {}
    merged: dict[str, str] = {}
    for src in (payload_data, context_data):
        for key, value in src.items():
            k = str(key or "").strip()
            if not k:
                continue
            merged[k] = normalize_space(str(value or ""))

    merged["yyyymmdd##"] = _resolve_yyyymmdd_serial_no(context_data, payload_data)
    merged["#"] = merged["yyyymmdd##"]
    merged["yyyy-MM-dd"] = _resolve_iso_date_text(context_data, payload_data)
    merged["selected_rows"] = normalize_space(
        str(
            context_data.get("selected_rows")
            or context_data.get("cylinder_total_count")
            or payload_data.get("selected_rows")
            or payload_data.get("cylinder_total_count")
            or "",
        ),
    )
    merged["ownership_code"] = normalize_space(
        str(
            context_data.get("ownership_code")
            or context_data.get("report_owner_name")
            or context_data.get("owner_code")
            or payload_data.get("ownership_code")
            or payload_data.get("report_owner_name")
            or payload_data.get("owner_code")
            or "",
        ),
    )
    merged["filling_medium"] = normalize_space(
        str(
            context_data.get("filling_medium")
            or context_data.get("gas_type")
            or context_data.get("medium")
            or payload_data.get("filling_medium")
            or payload_data.get("gas_type")
            or payload_data.get("medium")
            or "",
        ),
    )
    return merged


def _fill_jinja_placeholders_in_root(root: ET.Element, values: dict[str, str]) -> bool:
    changed = False
    valid_values = {str(k or "").strip(): str(v or "") for k, v in (values or {}).items() if str(k or "").strip()}
    if not valid_values:
        return False
    token_pattern = re.compile(r"{{\s*([A-Za-z0-9_#-]+)\s*}}")
    unresolved_pattern = re.compile(r"{{[^{}]{1,120}}}")
    for node in root.findall(".//w:t", NS):
        text = str(node.text or "")
        if "{{" not in text or "}}" not in text:
            continue
        replaced = token_pattern.sub(lambda m: valid_values.get(str(m.group(1) or "").strip(), ""), text)
        replaced = unresolved_pattern.sub("", replaced)
        if replaced != text:
            node.text = replaced
            changed = True
    for paragraph in root.findall(".//w:p", NS):
        text_nodes = paragraph.findall(".//w:t", NS)
        if not text_nodes:
            continue
        if paragraph.findall(".//w:br", NS):
            continue
        full_text = "".join(str(node.text or "") for node in text_nodes)
        if "{{" not in full_text or "}}" not in full_text:
            continue
        replaced = token_pattern.sub(lambda m: valid_values.get(str(m.group(1) or "").strip(), ""), full_text)
        replaced = unresolved_pattern.sub("", replaced)
        if replaced == full_text:
            continue
        text_nodes[0].text = replaced
        for node in text_nodes[1:]:
            node.text = ""
        changed = True
    return changed


def _build_placeholder_values_from_rules(
    template_name: str,
    context: dict[str, Any],
    payload: dict[str, Any],
) -> dict[str, str]:
    rules = get_fill_placeholders(template_name)
    if not rules:
        return {}
    context_data = context if isinstance(context, dict) else {}
    payload_data = payload if isinstance(payload, dict) else {}

    def read_value(key: str) -> str:
        normalized_key = str(key or "").strip()
        if not normalized_key:
            return ""
        value = context_data.get(normalized_key, "")
        if value is None or str(value).strip() == "":
            value = payload_data.get(normalized_key, "")
        return normalize_space(str(value or ""))

    result: dict[str, str] = {}
    for row in rules:
        marker = str((row or {}).get("marker", "")).strip()
        key = str((row or {}).get("key", "")).strip()
        if not marker or not key:
            continue
        value = read_value(key)
        if not value:
            for fallback_key in (row.get("fallback_keys") or []):
                value = read_value(str(fallback_key or ""))
                if value:
                    break
        if not value:
            value = normalize_space(str(row.get("default", "") or ""))
        result[marker] = value
    return result


def _parse_appendix1_rows_text(value: str) -> list[tuple[str, str, str]]:
    text = str(value or "").replace("\r\n", "\n").replace("\r", "\n").strip()
    if not text:
        return []
    result: list[tuple[str, str, str]] = []
    for raw_line in text.split("\n"):
        line = str(raw_line or "").strip()
        if not line:
            continue
        if "\t" in line:
            parts = [x.strip() for x in line.split("\t")]
        elif "," in line:
            parts = [x.strip() for x in line.split(",")]
        else:
            parts = [x.strip() for x in re.split(r"\s{2,}", line)]
        parts = [x for x in parts if x]
        if len(parts) < 3:
            continue
        cylinder_no = str(parts[0] or "").strip()
        maker_code = str(parts[1] or "").strip().upper()
        next_date = str(parts[2] or "").strip()
        result.append((cylinder_no, maker_code, next_date))
    return result


def _find_appendix1_table(tables: list[ET.Element]) -> ET.Element | None:
    for table in tables:
        table_text = " ".join(get_cell_text(tc) for tc in table.findall(".//w:tc", NS))
        if "气瓶编号" in table_text and "制造单位名称或代号" in table_text and "下次检验日期" in table_text:
            return table
    return None


def _find_appendix1_header_row_index(rows: list[ET.Element]) -> int:
    for idx, row in enumerate(rows):
        cells = row.findall("./w:tc", NS)
        row_text = " ".join(get_cell_text(cell) for cell in cells)
        if "气瓶编号" in row_text and "制造单位名称或代号" in row_text and "下次检验日期" in row_text:
            return idx
    return -1


def _fill_appendix1_table_rows_from_context(root: ET.Element, context: dict[str, Any]) -> bool:
    rows_data = _parse_appendix1_rows_text(str((context or {}).get("appendix1_rows_text", "")))
    if not rows_data:
        return False
    tables = root.findall(".//w:tbl", NS)
    target_table = _find_appendix1_table(tables)
    if target_table is None:
        return False
    rows = target_table.findall("./w:tr", NS)
    if not rows:
        return False
    header_idx = _find_appendix1_header_row_index(rows)
    if header_idx < 0:
        return False
    data_rows = rows[header_idx + 1:]
    if not data_rows:
        return False

    changed = False
    if len(data_rows) < len(rows_data):
        template_row = data_rows[0]
        for _ in range(len(rows_data) - len(data_rows)):
            target_table.append(deepcopy(template_row))
        rows = target_table.findall("./w:tr", NS)
        data_rows = rows[header_idx + 1:]
        changed = True

    effective_rows = data_rows[:len(rows_data)]
    for row_index, row in enumerate(effective_rows):
        cells = row.findall("./w:tc", NS)
        if len(cells) < 4:
            continue
        cylinder_no, maker_code, next_date = rows_data[row_index]
        set_cell_text(cells[0], str(row_index + 1))
        set_cell_text(cells[1], cylinder_no)
        set_cell_text(cells[2], maker_code)
        set_cell_text(cells[3], next_date)
        changed = True

    # Remove extra blank template rows to keep appendix aligned with sample format.
    if len(data_rows) > len(rows_data):
        for row in data_rows[len(rows_data):]:
            try:
                target_table.remove(row)
                changed = True
            except Exception:
                continue
    return changed


def _iter_docx_paragraphs_for_signatures(document: Any) -> list[Any]:
    result: list[Any] = []
    for paragraph in list(getattr(document, "paragraphs", []) or []):
        result.append(paragraph)

    def walk_tables(tables: list[Any]) -> None:
        for table in tables or []:
            for row in getattr(table, "rows", []) or []:
                for cell in getattr(row, "cells", []) or []:
                    for paragraph in getattr(cell, "paragraphs", []) or []:
                        result.append(paragraph)
                    walk_tables(getattr(cell, "tables", []) or [])

    walk_tables(list(getattr(document, "tables", []) or []))
    return result


def _clear_docx_paragraph_runs(paragraph: Any) -> None:
    runs = list(getattr(paragraph, "runs", []) or [])
    for run in runs:
        try:
            run._element.getparent().remove(run._element)
        except Exception:
            continue


def _fill_signature_images_for_output_docx(output_path: Path, payload: dict[str, Any]) -> bool:
    if Document is None or Mm is None:
        return False
    placeholder_to_value = {
        "【检验员签字图片】": str(payload.get("inspector_sign_image", "")).strip(),
        "【审核签字图片】": str(payload.get("reviewer_sign_image", "")).strip(),
        "【批准签字图片】": str(payload.get("approver_sign_image", "")).strip(),
    }
    placeholder_to_path: dict[str, Path] = {}
    for placeholder, value in placeholder_to_value.items():
        if not value:
            continue
        image_path = resolve_signature_image_path(value)
        if image_path is None:
            continue
        placeholder_to_path[placeholder] = image_path
    if not placeholder_to_path:
        return False

    try:
        doc = Document(str(output_path))
    except Exception:
        return False

    updated = False
    placeholders = sorted(placeholder_to_path.keys(), key=lambda x: len(x), reverse=True)
    pattern = re.compile("|".join(re.escape(x) for x in placeholders))

    for paragraph in _iter_docx_paragraphs_for_signatures(doc):
        text = str(getattr(paragraph, "text", "") or "")
        if not text:
            continue
        matches = list(pattern.finditer(text))
        if not matches:
            continue
        _clear_docx_paragraph_runs(paragraph)
        cursor = 0
        for match in matches:
            start, end = match.span()
            if start > cursor:
                paragraph.add_run(text[cursor:start])
            token = match.group(0)
            image_path = placeholder_to_path.get(token)
            if image_path is not None:
                try:
                    run = paragraph.add_run()
                    run.add_picture(str(image_path), height=Mm(8))
                except Exception:
                    paragraph.add_run(token)
            else:
                paragraph.add_run(token)
            cursor = end
        if cursor < len(text):
            paragraph.add_run(text[cursor:])
        updated = True

    if not updated:
        return False
    try:
        doc.save(str(output_path))
        return True
    except Exception:
        return False


def _fill_header_base_fields_xml(xml_data: bytes, payload: dict[str, Any]) -> bytes:
    try:
        original_namespaces = _capture_namespaces(xml_data)
        root = ET.fromstring(xml_data)
    except Exception:
        return xml_data

    tables = root.findall(".//w:tbl", NS)
    if tables:
        _fill_generic_base_labels_in_tables(tables, payload)
    _fill_generic_base_labels_in_paragraphs(root, payload)
    _preserve_original_namespaces(root, original_namespaces)
    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def _fill_r803b_hammer_actual_rows(tbl: ET.Element, rows_data: list[list[str]]) -> None:
    rows = tbl.findall("./w:tr", NS)
    actual_rows: list[list[ET.Element]] = []
    for tr in rows:
        cells = tr.findall("./w:tc", NS)
        if not cells:
            continue
        first_text = get_cell_text(cells[0])
        if "实际值(g)" in first_text:
            actual_rows.append(cells)

    if not actual_rows:
        return

    for cells in actual_rows:
        for idx in range(1, len(cells)):
            set_cell_text(cells[idx], "")

    for row_idx, values in enumerate(rows_data):
        if row_idx >= len(actual_rows):
            break
        cells = actual_rows[row_idx]
        for idx, value in enumerate(values, start=1):
            if idx >= len(cells):
                break
            set_cell_text(cells[idx], normalize_space(value))


def _fill_r803b_result_checks(tbl: ET.Element) -> None:
    for tr in tbl.findall("./w:tr", NS):
        cells = tr.findall("./w:tc", NS)
        if not cells:
            continue
        for cell in cells:
            compact = re.sub(r"\s+", "", get_cell_text(cell))
            if not compact.startswith("结果"):
                continue
            set_cell_text(cell, "结果：√")


def _fill_generic_result_checks_in_tables(tables: list[ET.Element]) -> bool:
    changed = False
    for tbl in tables:
        for tr in tbl.findall("./w:tr", NS):
            cells = tr.findall("./w:tc", NS)
            if not cells:
                continue
            for cell in cells:
                compact = re.sub(r"\s+", "", get_cell_text(cell))
                if not compact.startswith("结果"):
                    continue
                if any(mark in compact for mark in ("■", "☑", "√")):
                    continue
                set_cell_text(cell, "结果：√")
                changed = True
    return changed


def _fill_generic_semantic_value_matrices_from_payload(tables: list[ET.Element], payload: dict[str, Any]) -> bool:
    detail_text = str(payload.get("detail_general_check", "") or "")
    source_maps = build_semantic_value_maps_from_general_check_text(detail_text, normalize_space=normalize_space)
    if not source_maps:
        return False

    changed = False
    for tbl in tables:
        rows = tbl.findall("./w:tr", NS)
        if not rows:
            continue
        header_row_idx = -1
        key_col_idx = -1
        value_col_idx = -1
        for ri, tr in enumerate(rows):
            cells = tr.findall("./w:tc", NS)
            if not cells:
                continue
            texts = [normalize_space(get_cell_text(cell)) for cell in cells]
            key_col_idx, value_col_idx = detect_semantic_key_value_columns(texts)
            if key_col_idx >= 0 and value_col_idx >= 0:
                header_row_idx = ri
                break
        if header_row_idx < 0 or key_col_idx < 0 or value_col_idx < 0:
            continue

        key_col_carry = ""
        for ri in range(header_row_idx + 1, len(rows)):
            cells = rows[ri].findall("./w:tc", NS)
            if not cells:
                continue
            if key_col_idx >= len(cells) or value_col_idx >= len(cells):
                continue
            key_text = normalize_space(get_cell_text(cells[key_col_idx]))
            if key_text:
                key_col_carry = key_text
            else:
                key_text = key_col_carry
            if not key_text:
                continue
            key = normalize_semantic_key(key_text, normalize_space=normalize_space)
            if not key:
                continue
            fill_value = pick_semantic_value_for_key(source_maps, key, normalize_space=normalize_space)
            if not fill_value:
                continue
            current_value = normalize_space(get_cell_text(cells[value_col_idx]))
            if current_value == fill_value:
                continue
            set_cell_text(cells[value_col_idx], fill_value)
            changed = True
    return changed


def _fill_generic_uncertainty_u_from_payload(root: ET.Element, tables: list[ET.Element], payload: dict[str, Any]) -> bool:
    detail_text = str(payload.get("detail_general_check", "") or "")
    uncertainty_items = extract_uncertainty_items(detail_text, normalize_space=normalize_space)
    measured_items = extract_measured_value_items(detail_text, normalize_space=normalize_space)
    u_value = extract_uncertainty_u_value(detail_text, normalize_space=normalize_space)
    if not u_value and not uncertainty_items and not measured_items:
        raw_text = normalize_multiline_text(str(payload.get("__raw_record_text", "") or ""), normalize_space=normalize_space)
        if raw_text:
            uncertainty_items = extract_uncertainty_items(raw_text, normalize_space=normalize_space)
            measured_items = extract_measured_value_items(raw_text, normalize_space=normalize_space)
            u_value = extract_uncertainty_u_value(raw_text, normalize_space=normalize_space)
    if not u_value and not uncertainty_items and not measured_items:
        return False
    changed = False
    paragraph_hits = 0
    table_hits = 0
    paragraph_anchor_hint = ""
    for paragraph in root.findall(".//w:p", NS):
        current = normalize_space("".join([(node.text or "") for node in paragraph.findall(".//w:t", NS)]))
        if not current:
            continue
        updated = replace_uncertainty_u_placeholder_by_items(current, uncertainty_items, normalize_space=normalize_space)
        if updated == current and u_value:
            updated = replace_uncertainty_u_placeholder(current, u_value, normalize_space=normalize_space)
        updated = replace_measured_value_placeholder_by_items(
            updated,
            measured_items,
            normalize_space=normalize_space,
            anchor_hint=paragraph_anchor_hint,
        )
        if updated != current:
            _set_paragraph_text(paragraph, updated)
            changed = True
            paragraph_hits += 1
        if "扩展不确定度" in current:
            paragraph_anchor_hint = current
    for tbl in tables:
        table_anchor_hint = ""
        for tr in tbl.findall("./w:tr", NS):
            for tc in tr.findall("./w:tc", NS):
                current = get_cell_text(tc)
                if not current:
                    continue
                updated = replace_uncertainty_u_placeholder_by_items(current, uncertainty_items, normalize_space=normalize_space)
                if updated == current and u_value:
                    updated = replace_uncertainty_u_placeholder(current, u_value, normalize_space=normalize_space)
                updated = replace_measured_value_placeholder_by_items(
                    updated,
                    measured_items,
                    normalize_space=normalize_space,
                    anchor_hint=table_anchor_hint,
                )
                if updated != current:
                    set_cell_text(tc, updated)
                    changed = True
                    table_hits += 1
                normalized_current = normalize_space(current)
                if "扩展不确定度" in normalized_current:
                    table_anchor_hint = normalized_current
    if logger.isEnabledFor(logging.DEBUG):
        logger.debug(
            "semantic_fill uncertainty/measured: u_value=%s uncertainty_items=%d measured_items=%d paragraph_hits=%d table_hits=%d",
            bool(u_value),
            len(uncertainty_items),
            len(measured_items),
            paragraph_hits,
            table_hits,
        )
    return changed


def _fill_generic_semantic_series_rows_from_payload(tables: list[ET.Element], payload: dict[str, Any]) -> bool:
    detail_text = str(payload.get("detail_general_check", "") or "")
    source_maps = build_series_row_value_maps_from_general_check_text(detail_text, normalize_space=normalize_space)
    if not source_maps:
        raw_text = normalize_multiline_text(str(payload.get("__raw_record_text", "") or ""), normalize_space=normalize_space)
        if raw_text:
            source_maps = build_series_row_value_maps_from_general_check_text(raw_text, normalize_space=normalize_space)
    if not source_maps:
        return False

    changed = False
    for tbl in tables:
        for tr in tbl.findall("./w:tr", NS):
            cells = tr.findall("./w:tc", NS)
            if len(cells) < 2:
                continue
            label = normalize_space(get_cell_text(cells[0]))
            if not label:
                continue
            source_values = pick_series_row_values_for_label(source_maps, label, normalize_space=normalize_space)
            if not source_values:
                continue
            for idx in range(1, min(len(cells), len(source_values) + 1)):
                fill_value = normalize_space(source_values[idx - 1])
                if not fill_value:
                    continue
                current = normalize_space(get_cell_text(cells[idx]))
                if current and current not in _PLACEHOLDER_VALUES:
                    continue
                if current == fill_value:
                    continue
                set_cell_text(cells[idx], fill_value)
                changed = True
    return changed


def _is_r882_profile(payload: dict[str, Any]) -> bool:
    name = str(payload.get("__template_name", "") or "")
    return bool(re.search(r"r[-_ ]?882b|屏蔽室", name, flags=re.IGNORECASE))


def _extract_r882_background_noise_values(text: str) -> list[str]:
    lines = [normalize_space(x) for x in str(text or "").splitlines() if normalize_space(x)]
    if not lines:
        return []
    bg0 = ""
    working: list[str] = []
    for line in lines:
        if "背景噪声" not in line:
            continue
        nums = re.findall(r"([+-]?\d+(?:\.\d+)?)\s*pC", line, flags=re.IGNORECASE)
        if not nums:
            continue
        value = normalize_space(nums[-1])
        if not value:
            continue
        if re.search(r"(?<!\d)0\s*kV", line, flags=re.IGNORECASE):
            if not bg0:
                bg0 = value
        else:
            working.append(value)
    out: list[str] = []
    if bg0:
        out.append(bg0)
    out.extend([x for x in working if x])
    return out


def _extract_r882_series_rows_from_text(text: str) -> list[tuple[str, str, str]]:
    rows = [x for x in str(text or "").splitlines() if str(x or "").strip()]
    if not rows:
        return []

    grid: list[list[str]] = []
    for row in rows:
        if "\t" not in row:
            continue
        grid.append([normalize_space(x) for x in row.split("\t")])
    if not grid:
        return []

    out: list[tuple[str, str, str]] = []
    for i, row in enumerate(grid):
        p1_idx = -1
        p2_idx = -1
        se_idx = -1
        for ci, cell in enumerate(row):
            compact = re.sub(r"\s+", "", str(cell or ""))
            if p1_idx < 0 and "P1" in compact:
                p1_idx = ci
            if p2_idx < 0 and "P2" in compact:
                p2_idx = ci
            if se_idx < 0 and ("SE" in compact or "屏蔽效能" in compact):
                se_idx = ci
        if p1_idx < 0 or p2_idx < 0:
            continue

        for j in range(i + 1, len(grid)):
            data = grid[j]
            if len(data) <= max(p1_idx, p2_idx):
                continue
            row_text = normalize_space(" ".join(data))
            if not row_text:
                continue
            if "P1" in row_text and "P2" in row_text:
                break
            p1 = _extract_first_number(data[p1_idx])
            p2 = _extract_first_number(data[p2_idx])
            se = _extract_first_number(data[se_idx]) if se_idx >= 0 and se_idx < len(data) else ""
            if not p1 and not p2 and not se:
                continue
            if not se and p1 and p2:
                try:
                    se = _format_decimal(float(p1) - float(p2))
                except Exception:
                    se = ""
            out.append((p1, p2, se))
    return out


def _extract_first_number(value: str) -> str:
    m = re.search(r"[+-]?\d+(?:\.\d+)?", str(value or ""))
    if not m:
        return ""
    return normalize_space(m.group(0))


def _format_decimal(value: float) -> str:
    text = f"{value:.3f}".rstrip("0").rstrip(".")
    return text if text else "0"


def _is_numeric_placeholder_cell(value: str) -> bool:
    text = normalize_space(value)
    if not text:
        return True
    if re.search(r"[+-]?\d+(?:\.\d+)?", text):
        return False
    return bool(re.fullmatch(r"[—\-_/\.…\s]*", text))


def _fill_r882b_background_noise_placeholders(root: ET.Element, values: list[str]) -> bool:
    if not values:
        return False

    changed = False
    idx = 0

    def next_value() -> str:
        nonlocal idx
        if idx < len(values):
            value = values[idx]
            idx += 1
            return value
        return values[-1]

    def replace_text(source: str) -> str:
        if "背景噪声" not in source or "pC" not in source:
            return source
        pattern = re.compile(r"(背景噪声[^。\n]*[:：]\s*)([^。\n]*?)(\s*pC)", flags=re.IGNORECASE)

        def repl(match: re.Match[str]) -> str:
            current = normalize_space(match.group(2))
            if re.search(r"[+-]?\d+(?:\.\d+)?", current):
                return match.group(0)
            value = next_value()
            return f"{match.group(1)}{value}{match.group(3)}"

        return pattern.sub(repl, source)

    for paragraph in root.findall(".//w:p", NS):
        current = normalize_space("".join([(node.text or "") for node in paragraph.findall(".//w:t", NS)]))
        if not current:
            continue
        updated = replace_text(current)
        if updated != current:
            _set_paragraph_text(paragraph, updated)
            changed = True

    for tc in root.findall(".//w:tc", NS):
        current = get_cell_text(tc)
        if not current:
            continue
        updated = replace_text(current)
        if updated != current:
            set_cell_text(tc, updated)
            changed = True
    return changed


def _fill_r882b_p_series_tables(tables: list[ET.Element], rows_data: list[tuple[str, str, str]]) -> bool:
    if not rows_data:
        return False

    changed = False
    cursor = 0
    for tbl in tables:
        rows = tbl.findall("./w:tr", NS)
        if not rows:
            continue
        p1_idx = -1
        p2_idx = -1
        se_idx = -1
        header_row_idx = -1
        for ri, tr in enumerate(rows):
            cells = tr.findall("./w:tc", NS)
            texts = [normalize_space(get_cell_text(tc)) for tc in cells]
            for ci, text in enumerate(texts):
                compact = re.sub(r"\s+", "", text)
                if p1_idx < 0 and "P1" in compact:
                    p1_idx = ci
                if p2_idx < 0 and "P2" in compact:
                    p2_idx = ci
                if se_idx < 0 and ("SE" in compact or "屏蔽效能" in compact):
                    se_idx = ci
            if p1_idx >= 0 and p2_idx >= 0:
                header_row_idx = ri
                break
        if header_row_idx < 0:
            continue
        if se_idx < 0:
            se_idx = p2_idx + 1

        for ri in range(header_row_idx + 1, len(rows)):
            if cursor >= len(rows_data):
                break
            tr = rows[ri]
            cells = tr.findall("./w:tc", NS)
            if len(cells) <= max(p1_idx, p2_idx, se_idx):
                continue
            p1, p2, se = rows_data[cursor]
            current_p1 = get_cell_text(cells[p1_idx])
            current_p2 = get_cell_text(cells[p2_idx])
            current_se = get_cell_text(cells[se_idx]) if se_idx < len(cells) else ""
            wrote = False
            if p1 and _is_numeric_placeholder_cell(current_p1):
                set_cell_text(cells[p1_idx], p1)
                wrote = True
            if p2 and _is_numeric_placeholder_cell(current_p2):
                set_cell_text(cells[p2_idx], p2)
                wrote = True
            if se and se_idx < len(cells) and _is_numeric_placeholder_cell(current_se):
                set_cell_text(cells[se_idx], se)
                wrote = True
            if wrote:
                changed = True
                cursor += 1
    return changed


def _fill_r882b_specific_sections(
    root: ET.Element,
    tables: list[ET.Element],
    payload: dict[str, Any],
) -> bool:
    if not _is_r882_profile(payload):
        return False

    detail_text = normalize_multiline_text_preserve_tabs(str(payload.get("detail_general_check", "") or ""), normalize_space=normalize_space)
    measurement_items_text = normalize_multiline_text_preserve_tabs(str(payload.get("__measurement_items_text", "") or ""), normalize_space=normalize_space)
    raw_text = normalize_multiline_text(str(payload.get("__raw_record_text", "") or ""), normalize_space=normalize_space)
    combined = "\n".join([x for x in (detail_text, measurement_items_text, raw_text) if x])

    background_values = _extract_r882_background_noise_values(combined)
    if payload.get("shield_background_noise_0kv_pc"):
        bg0 = normalize_space(str(payload.get("shield_background_noise_0kv_pc", "")))
        if bg0:
            if background_values:
                background_values[0] = bg0
            else:
                background_values = [bg0]
    if payload.get("shield_background_noise_working_kv_pc"):
        bgw = normalize_space(str(payload.get("shield_background_noise_working_kv_pc", "")))
        if bgw:
            if len(background_values) >= 2:
                background_values[1] = bgw
            elif len(background_values) == 1:
                background_values.append(bgw)
            else:
                background_values = [bgw]

    rows_data = _extract_r882_series_rows_from_text(measurement_items_text or detail_text or raw_text)
    changed = False
    changed = _fill_r882b_background_noise_placeholders(root, background_values) or changed
    changed = _fill_r882b_p_series_tables(tables, rows_data) or changed
    return changed


def _fill_generic_result_checks_by_semantics(tables: list[ET.Element], payload: dict[str, Any]) -> bool:
    source_lines = extract_source_general_check_lines(str(payload.get("detail_general_check", "")))
    if not source_lines:
        return False
    if _is_r872_profile(payload):
        return _fill_r872_result_checks_by_rules(tables, source_lines)

    changed = False
    marked_count = 0
    unmarked_count = 0
    used_source_indexes: set[int] = set()
    for tbl in tables:
        rows = tbl.findall("./w:tr", NS)
        previous_requirement = ""
        for tr in rows:
            cells = tr.findall("./w:tc", NS)
            if not cells:
                continue
            result_cells = [cell for cell in cells if re.sub(r"\s+", "", get_cell_text(cell)).startswith("结果")]
            row_requirement = _extract_requirement_text_from_cells(cells)
            if not result_cells:
                if row_requirement:
                    previous_requirement = row_requirement
                continue

            target_requirement = row_requirement or previous_requirement
            matched_index, _ = match_best_source_line(
                target_text=target_requirement,
                source_lines=source_lines,
                used_indexes=used_source_indexes,
                threshold=0.30,
            )
            if matched_index >= 0:
                source_line = source_lines[matched_index]
                if not is_reliable_result_semantic_match(target_requirement, source_line, normalize_space=normalize_space):
                    matched_index = -1
            for result_cell in result_cells:
                if matched_index >= 0:
                    set_cell_text(result_cell, "结果：√")
                    changed = True
                    marked_count += 1
                else:
                    set_cell_text(result_cell, "结果：")
                    changed = True
                    unmarked_count += 1
            if matched_index >= 0:
                used_source_indexes.add(matched_index)
    if logger.isEnabledFor(logging.DEBUG):
        logger.debug(
            "semantic_fill result_checks: source_lines=%d marked=%d unmarked=%d matched_sources=%d",
            len(source_lines),
            marked_count,
            unmarked_count,
            len(used_source_indexes),
        )
    return changed



def _fill_r872_result_checks_by_rules(tables: list[ET.Element], source_lines: list[str]) -> bool:
    changed = False
    for tbl in tables:
        rows = tbl.findall("./w:tr", NS)
        for tr in rows:
            cells = tr.findall("./w:tc", NS)
            if not cells:
                continue
            result_cells = [cell for cell in cells if re.sub(r"\s+", "", get_cell_text(cell)).startswith("结果")]
            non_result_cells = [cell for cell in cells if cell not in result_cells]
            for cell in non_result_cells:
                original_text = get_cell_text(cell)
                updated_text = fill_r872_requirement_text(original_text, source_lines)
                if updated_text != original_text:
                    set_cell_text(cell, updated_text)
                    changed = True
            row_requirement = _extract_requirement_text_from_cells(cells)
            if not result_cells:
                continue
            # R872按“当前结果行的要求文本”判定，不跨行继承，避免错位误勾。
            mark = should_mark_r872_result(target_requirement=row_requirement, source_lines=source_lines)
            for result_cell in result_cells:
                set_cell_text(result_cell, "结果：√" if mark else "结果：")
                changed = True
    return changed


def _is_r872_profile(payload: dict[str, Any]) -> bool:
    template_name = normalize_space(str(payload.get("__template_name", ""))).lower()
    if re.search(r"r[-_ ]?872b", template_name, flags=re.IGNORECASE):
        return True
    device_name = normalize_space(str(payload.get("device_name", "")))
    return bool(re.search(r"扭转", device_name))


def _extract_requirement_text_from_cells(cells: list[ET.Element]) -> str:
    parts: list[str] = []
    for cell in cells:
        text = normalize_space(get_cell_text(cell))
        if not text:
            continue
        compact = re.sub(r"\s+", "", text)
        if compact.startswith("结果"):
            continue
        if re.search(r"^(?:[一二三四五六七八九十]+[、.．)]|一般检查[:：]?)$", text):
            continue
        parts.append(text)
    return normalize_space(" ".join(parts))


def _should_auto_fill_result_checks(template_path: Path, payload: dict[str, Any]) -> bool:
    text = normalize_space(str(payload.get("auto_check_results", "")))
    if text in {"0", "false", "no", "off"}:
        return False
    name = normalize_space(template_path.name).lower()
    return bool(re.search(r"r[-_ ]?859b", name, flags=re.IGNORECASE))


def _should_semantic_fill_result_checks(template_path: Path) -> bool:
    name = normalize_space(template_path.name).lower()
    return bool(re.search(r"r[-_ ]?872b|扭转", name, flags=re.IGNORECASE))


def _fill_page_number_placeholder(tables: list[ET.Element]) -> None:
    # Kept for backward compatibility, unified flow uses _fill_page_number_placeholders_in_root.
    for tbl in tables:
        for tr in tbl.findall("./w:tr", NS):
            cells = tr.findall("./w:tc", NS)
            for cell in cells:
                text = normalize_space(get_cell_text(cell))
                if not _is_page_placeholder_text(text):
                    continue
                set_cell_page_fields(cell, 1, 1)


def _fill_page_number_placeholder_in_paragraphs(root: ET.Element) -> None:
    # Kept for backward compatibility, unified flow uses _fill_page_number_placeholders_in_root.
    for paragraph in root.findall(".//w:p", NS):
        text = normalize_space("".join([(node.text or "") for node in paragraph.findall(".//w:t", NS)]))
        if not _is_page_placeholder_text(text):
            continue
        _set_paragraph_text(paragraph, "第 1 页/共 1 页")


def _append_simple_field(paragraph: ET.Element, instr_name: str) -> None:
    field = ET.SubElement(paragraph, f"{{{W_NS}}}fldSimple")
    field.set(f"{{{W_NS}}}instr", f"{instr_name} \\* MERGEFORMAT")
    run = ET.SubElement(field, f"{{{W_NS}}}r")
    run_props = ET.SubElement(run, f"{{{W_NS}}}rPr")
    ET.SubElement(run_props, f"{{{W_NS}}}noProof")
    text = ET.SubElement(run, f"{{{W_NS}}}t")
    text.text = "1"


def _extract_location_from_other_calibration_info(text: str) -> str:
    return _extract_location_from_other_calibration_info_bridge(text, extract_value_by_regex)


def _extract_temperature_from_other_calibration_info(text: str) -> str:
    return _extract_temperature_from_other_calibration_info_bridge(text, extract_value_by_regex)


def _extract_humidity_from_other_calibration_info(text: str) -> str:
    return _extract_humidity_from_other_calibration_info_bridge(text, extract_value_by_regex)


def _sanitize_location_text(value: str) -> str:
    return _sanitize_location_text_bridge(value)


def _replace_uncertainty_value(text: str, value: str, unit: str) -> str:
    return _replace_uncertainty_value_bridge(text, value, unit)


def _replace_measured_value(text: str, value: str, unit: str) -> str:
    return _replace_measured_value_bridge(text, value, unit)


def _find_cell_index_contains_any(cells: list[ET.Element], markers: tuple[str, ...]) -> int:
    return find_cell_index_contains_any(cells=cells, markers=markers, get_cell_text=get_cell_text)


def _extract_section_uncertainty(text: str, section_title: str, unit: str) -> str:
    return _extract_section_uncertainty_bridge(text, section_title, unit, extract_value_by_regex)


def _extract_section_measured_value(text: str, section_title: str, unit: str) -> str:
    return _extract_section_measured_value_bridge(text, section_title, unit, extract_value_by_regex)
